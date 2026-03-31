import sys
import os
import json
from datetime import datetime
import numpy as np
import tempfile
import subprocess
import requests
import asyncio
from PyQt5.QtCore import QThread, pyqtSignal, QTimer, Qt, QSize, QDate
from PyQt5.QtWidgets import (
    QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout,
    QLabel, QPushButton, QLineEdit, QListWidget, QListWidgetItem,
    QTabWidget, QFileDialog, QMessageBox, QScrollArea, QCheckBox,
    QGroupBox, QGridLayout, QFrame, QComboBox, QTableWidget, QTableWidgetItem,
    QHeaderView, QTextEdit, QAbstractItemView, QDialog, QFormLayout,
    QDialogButtonBox, QCalendarWidget, QRadioButton, QButtonGroup, QProgressDialog
)
from PyQt5.QtGui import QPixmap, QIcon, QFont

# Módulos do sistema
from analytics_service import (
    init_db, listar_erros, salvar_relatorio, listar_relatorios,
    obter_detalhes_relatorio, excluir_relatorio, listar_datas_disponiveis,
    listar_equipes_disponiveis, atualizar_status_relatorio,
    listar_status_disponiveis, sincronizar_erros_padrao, atualizar_nome_equipe
)
from matplotlib.backends.backend_qt5agg import FigureCanvasQTAgg as Canvas
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Cm, Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import pandas as pd
from openpyxl import load_workbook
from openpyxl.utils import get_column_letter
import re
from exportador_completo import exportar_dados_para_excel_fixo
from updater import GitHubUpdater

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
CONFIG_FILE = os.path.join(BASE_DIR, "config.json")

def resource_path(relative_path):
    try:
        base_path = sys._MEIPASS
    except AttributeError:
        base_path = os.path.abspath(".")
    return os.path.join(base_path, relative_path)

def carregar_configuracao():
    config_padrao = {
        "modelo_docx": "",
        "ultimo_diretorio": BASE_DIR,
        "empresa_nome": "ENGELMIG ENERGIA LTDA",
        "pasta_raiz": "",
        "pasta_excel": "",
        "icone_app": ""
    }
    if os.path.exists(CONFIG_FILE):
        try:
            with open(CONFIG_FILE, 'r', encoding='utf-8') as f:
                return {**config_padrao, **json.load(f)}
        except:
            return config_padrao
    return config_padrao

def salvar_configuracao(config):
    try:
        with open(CONFIG_FILE, 'w', encoding='utf-8') as f:
            json.dump(config, f, indent=4, ensure_ascii=False)
        return True
    except:
        return False

def obter_versao_atual():
    try:
        version_path = resource_path("version.txt")
        with open(version_path, "r") as f:
            return f.read().strip()
    except:
        return "1.0.0"

class UpdateChecker(QThread):
    update_available = pyqtSignal(str)  # sinal com a URL de download
    finished = pyqtSignal()
    
    def __init__(self, repo, current_version, parent=None):
        super().__init__(parent)
        self.repo = repo
        self.current_version = current_version
    
    def run(self):
        updater = GitHubUpdater(self.repo, self.current_version)
        has_update, url = updater.check_and_update()
        if has_update:
            self.update_available.emit(url)
        self.finished.emit()

class UpdateDownloadThread(QThread):
    progress = pyqtSignal(int)
    finished = pyqtSignal(str, str)
    def __init__(self, url, destino):
        super().__init__()
        self.url = url
        self.destino = destino
    def run(self):
        try:
            response = requests.get(self.url, stream=True)
            total_size = int(response.headers.get('content-length', 0))
            downloaded = 0
            with open(self.destino, 'wb') as f:
                for chunk in response.iter_content(chunk_size=8192):
                    f.write(chunk)
                    downloaded += len(chunk)
                    if total_size:
                        percent = int(downloaded * 100 / total_size)
                        self.progress.emit(percent)
            self.finished.emit(self.destino, "")
        except Exception as e:
            self.finished.emit("", str(e))

class ExportThread(QThread):
    finished = pyqtSignal(object)
    def __init__(self, export_func):
        super().__init__()
        self.export_func = export_func
    def run(self):
        try:
            resultado = self.export_func()
            self.finished.emit(resultado)
        except Exception as e:
            self.finished.emit({'sucesso': False, 'erro': str(e)})

class ConfiguracaoDialog(QDialog):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setWindowTitle("Configurações do Sistema")
        self.setModal(True)
        self.resize(600, 500)
        self.config = carregar_configuracao()
        self.init_ui()

    def init_ui(self):
        layout = QVBoxLayout(self)
        titulo = QLabel("⚙️ CONFIGURAÇÕES DO SISTEMA")
        titulo.setStyleSheet("font-size:16px; font-weight:bold; color:#1976D2; padding:10px; background-color:#E3F2FD; border-radius:5px;")
        titulo.setAlignment(Qt.AlignCenter)
        layout.addWidget(titulo)

        grupo_modelo = QGroupBox("📄 Modelo de Relatório")
        layout_modelo = QVBoxLayout(grupo_modelo)
        self.lbl_caminho_valor = QLabel(self.config.get("modelo_docx", "Não configurado"))
        self.lbl_caminho_valor.setWordWrap(True)
        self.lbl_caminho_valor.setStyleSheet("background-color:#F5F5F5; padding:8px; border-radius:3px; border:1px solid #E0E0E0;")
        btn_selecionar = QPushButton("📁 Selecionar Novo Modelo")
        btn_selecionar.setStyleSheet("background-color:#4CAF50; color:white; padding:10px; font-weight:bold;")
        btn_selecionar.clicked.connect(self.selecionar_modelo)
        layout_modelo.addWidget(self.lbl_caminho_valor)
        layout_modelo.addWidget(btn_selecionar)
        layout.addWidget(grupo_modelo)

        grupo_icone = QGroupBox("🖼️ Ícone da Aplicação")
        layout_icone = QVBoxLayout(grupo_icone)
        self.lbl_icone_valor = QLabel(self.config.get("icone_app", "Não configurado"))
        self.lbl_icone_valor.setWordWrap(True)
        self.lbl_icone_valor.setStyleSheet("background-color:#F5F5F5; padding:8px; border-radius:3px; border:1px solid #E0E0E0;")
        self.lbl_preview_icone = QLabel()
        self.lbl_preview_icone.setFixedSize(64, 64)
        self.lbl_preview_icone.setStyleSheet("border:1px solid #BDBDBD; background-color:white;")
        self.lbl_preview_icone.setAlignment(Qt.AlignCenter)
        self.atualizar_preview_icone()
        btn_selecionar_icone = QPushButton("📁 Selecionar Ícone")
        btn_selecionar_icone.setStyleSheet("background-color:#9C27B0; color:white; padding:10px; font-weight:bold;")
        btn_selecionar_icone.clicked.connect(self.selecionar_icone)
        layout_icone.addWidget(self.lbl_icone_valor)
        layout_icone.addWidget(self.lbl_preview_icone)
        layout_icone.addWidget(btn_selecionar_icone)
        layout.addWidget(grupo_icone)

        button_box = QDialogButtonBox(QDialogButtonBox.Save | QDialogButtonBox.Cancel)
        button_box.accepted.connect(self.salvar_config)
        button_box.rejected.connect(self.reject)
        layout.addWidget(button_box)

    def selecionar_modelo(self):
        ultimo_dir = self.config.get("ultimo_diretorio", BASE_DIR)
        caminho, _ = QFileDialog.getOpenFileName(self, "Selecionar Modelo", ultimo_dir, "Word (*.docx);;Todos (*.*)")
        if caminho:
            self.config["modelo_docx"] = caminho
            self.config["ultimo_diretorio"] = os.path.dirname(caminho)
            self.lbl_caminho_valor.setText(caminho)

    def selecionar_icone(self):
        ultimo_dir = self.config.get("ultimo_diretorio", BASE_DIR)
        caminho, _ = QFileDialog.getOpenFileName(self, "Selecionar Ícone", ultimo_dir, "Imagens (*.ico *.png *.jpg);;Todos (*.*)")
        if caminho:
            self.config["icone_app"] = caminho
            self.config["ultimo_diretorio"] = os.path.dirname(caminho)
            self.lbl_icone_valor.setText(caminho)
            self.atualizar_preview_icone()

    def atualizar_preview_icone(self):
        icone_path = self.config.get("icone_app", "")
        if icone_path and os.path.exists(icone_path):
            pixmap = QPixmap(icone_path)
            if not pixmap.isNull():
                pixmap = pixmap.scaled(64, 64, Qt.KeepAspectRatio, Qt.SmoothTransformation)
                self.lbl_preview_icone.setPixmap(pixmap)
            else:
                self.lbl_preview_icone.setText("❌")
        else:
            self.lbl_preview_icone.setText("Sem ícone")

    def salvar_config(self):
        if salvar_configuracao(self.config):
            QMessageBox.information(self, "Configurações", "Salvas com sucesso!")
            if self.parent():
                self.parent().aplicar_icone_configurado()
            self.accept()

class ErroCheckBox(QCheckBox):
    def __init__(self, texto, categoria):
        super().__init__(texto)
        self.descricao = texto
        self.categoria = categoria
        self.setStyleSheet("""
            QCheckBox {
                padding: 8px;
                font-size: 11px;
            }
            QCheckBox:hover {
                background-color: #f0f0f0;
            }
            QCheckBox:checked {
                font-weight: bold;
                color: #2196F3;
            }
        """)

class AbaRelatorio(QWidget):
    def __init__(self, parent_app):
        super().__init__()
        self.parent_app = parent_app
        self.erros_selecionados = set()
        self.checkboxes = []
        self.categorias_expandidas = {}

        print(f"AbaRelatorio inicializada")
        
        layout_principal = QVBoxLayout(self)
        layout_principal.setSpacing(10)

        cabecalho = QLabel("📋 NOVO RELATÓRIO")
        cabecalho.setStyleSheet("""
            QLabel {
                font-size: 18px;
                font-weight: bold;
                color: #1976D2;
                padding: 10px;
                background-color: #E3F2FD;
                border-radius: 5px;
            }
        """)
        cabecalho.setAlignment(Qt.AlignCenter)
        layout_principal.addWidget(cabecalho)

        painel_info = QFrame()
        painel_info.setStyleSheet("""
            QFrame {
                background-color: #F5F5F5;
                border-radius: 5px;
                padding: 15px;
            }
        """)

        layout_info = QGridLayout(painel_info)
        layout_info.setSpacing(10)

        lbl_equipe = QLabel("👥 Equipe:")
        lbl_equipe.setStyleSheet("font-weight: bold; font-size: 14px;")
        self.equipe = QComboBox()
        self.equipe.setEditable(True)
        self.equipe.setPlaceholderText("Digite ou selecione a equipe...")
        self.equipe.setStyleSheet("""
            QComboBox {
                padding: 8px;
                border: 2px solid #BDBDBD;
                border-radius: 4px;
                font-size: 14px;
            }
            QComboBox:focus {
                border: 2px solid #2196F3;
            }
        """)
        self.carregar_equipes()

        lbl_data = QLabel("📅 Data:")
        lbl_data.setStyleSheet("font-weight: bold; font-size: 14px;")
        
        data_layout = QHBoxLayout()
        self.data = QLineEdit()
        self.data.setPlaceholderText("DD/MM/AAAA")
        self.data.setText(datetime.now().strftime("%d/%m/%Y"))
        self.data.setStyleSheet("""
            QLineEdit {
                padding: 8px;
                border: 2px solid #BDBDBD;
                border-radius: 4px;
                font-size: 14px;
            }
            QLineEdit:focus {
                border: 2px solid #2196F3;
            }
        """)

        btn_calendario = QPushButton("📅")
        btn_calendario.setToolTip("Selecionar data")
        btn_calendario.setFixedSize(40, 40)
        btn_calendario.setStyleSheet("""
            QPushButton {
                background-color: #4CAF50;
                color: white;
                font-size: 16px;
                border-radius: 4px;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #45a049;
            }
        """)
        btn_calendario.clicked.connect(self.abrir_calendario)

        data_layout.addWidget(self.data)
        data_layout.addWidget(btn_calendario)

        lbl_status = QLabel("📊 Status:")
        lbl_status.setStyleSheet("font-weight: bold; font-size: 14px;")
        
        status_layout = QHBoxLayout()
        self.status_group = QButtonGroup(self)
        
        self.btn_nao_conforme = QRadioButton("Não Conforme")
        self.btn_nao_conforme.setChecked(True)
        self.btn_nao_conforme.setStyleSheet("""
            QRadioButton {
                padding: 8px;
                font-size: 14px;
            }
            QRadioButton:checked {
                font-weight: bold;
                color: #D32F2F;
            }
        """)
        
        self.btn_conforme = QRadioButton("Conforme")
        self.btn_conforme.setStyleSheet("""
            QRadioButton {
                padding: 8px;
                font-size: 14px;
            }
            QRadioButton:checked {
                font-weight: bold;
                color: #388E3C;
            }
        """)
        
        self.status_group.addButton(self.btn_nao_conforme)
        self.status_group.addButton(self.btn_conforme)
        
        status_layout.addWidget(self.btn_nao_conforme)
        status_layout.addWidget(self.btn_conforme)
        status_layout.addStretch()

        layout_info.addWidget(lbl_equipe, 0, 0)
        layout_info.addWidget(self.equipe, 0, 1)
        layout_info.addWidget(lbl_data, 1, 0)
        layout_info.addLayout(data_layout, 1, 1)
        layout_info.addWidget(lbl_status, 2, 0)
        layout_info.addLayout(status_layout, 2, 1)

        layout_principal.addWidget(painel_info)

        btn_modelo = QPushButton("📄 Configurar Modelo de Papel Timbrado")
        btn_modelo.setStyleSheet("""
            QPushButton {
                background-color: #9C27B0;
                color: white;
                padding: 10px;
                font-weight: bold;
                border-radius: 5px;
                margin: 5px 0;
            }
            QPushButton:hover {
                background-color: #7B1FA2;
            }
        """)
        btn_modelo.clicked.connect(self.abrir_configuracao)
        layout_principal.addWidget(btn_modelo)

        self.contador_frame = QFrame()
        self.contador_frame.setStyleSheet("""
            QFrame {
                background-color: #FFF3E0;
                border: 2px solid #FF9800;
                border-radius: 5px;
                padding: 10px;
            }
        """)
        layout_contador = QHBoxLayout(self.contador_frame)

        self.lbl_contador = QLabel("✅ 0 erro(s) selecionado(s)")
        self.lbl_contador.setStyleSheet("""
            QLabel {
                font-weight: bold;
                color: #FF9800;
                font-size: 14px;
            }
        """)

        btn_limpar = QPushButton("🗑️ Limpar Seleção")
        btn_limpar.setStyleSheet("""
            QPushButton {
                background-color: #FFE0B2;
                color: #E65100;
                padding: 5px 10px;
                border-radius: 3px;
                font-size: 12px;
            }
            QPushButton:hover {
                background-color: #FFCC80;
            }
        """)
        btn_limpar.clicked.connect(self.limpar_selecao)

        layout_contador.addWidget(self.lbl_contador)
        layout_contador.addWidget(btn_limpar)
        layout_contador.addStretch()

        layout_principal.addWidget(self.contador_frame)

        painel_filtro = QFrame()
        painel_filtro.setStyleSheet("""
            QFrame {
                background-color: #F8F8F8;
                border: 1px solid #E0E0E0;
                border-radius: 5px;
                padding: 8px;
            }
        """)
        layout_filtro = QHBoxLayout(painel_filtro)

        lbl_filtro = QLabel("🔍 Filtrar por categoria:")
        lbl_filtro.setStyleSheet("font-weight: bold;")

        self.combo_filtro = QComboBox()
        self.combo_filtro.addItem("Todas as categorias")
        self.combo_filtro.setStyleSheet("""
            QComboBox {
                padding: 5px;
                border: 1px solid #BDBDBD;
                border-radius: 3px;
            }
        """)
        self.combo_filtro.currentTextChanged.connect(self.filtrar_erros)

        self.btn_expandir = QPushButton("📖 Expandir Tudo")
        self.btn_expandir.setStyleSheet("""
            QPushButton {
                background-color: #4CAF50;
                color: white;
                padding: 5px 10px;
                border-radius: 3px;
                font-size: 12px;
            }
            QPushButton:hover {
                background-color: #45a049;
            }
        """)
        self.btn_expandir.clicked.connect(self.expandir_tudo)

        self.btn_recolher = QPushButton("📕 Recolher Tudo")
        self.btn_recolher.setStyleSheet("""
            QPushButton {
                background-color: #757575;
                color: white;
                padding: 5px 10px;
                border-radius: 3px;
                font-size: 12px;
            }
            QPushButton:hover {
                background-color: #616161;
            }
        """)
        self.btn_recolher.clicked.connect(self.recolher_tudo)

        layout_filtro.addWidget(lbl_filtro)
        layout_filtro.addWidget(self.combo_filtro)
        layout_filtro.addStretch()
        layout_filtro.addWidget(self.btn_expandir)
        layout_filtro.addWidget(self.btn_recolher)

        layout_principal.addWidget(painel_filtro)

        self.scroll = QScrollArea()
        self.scroll.setWidgetResizable(True)
        self.scroll.setStyleSheet("QScrollArea { border: none; }")

        self.content_widget = QWidget()
        self.content_layout = QVBoxLayout(self.content_widget)
        self.content_layout.setSpacing(5)

        self.carregar_erros()

        self.scroll.setWidget(self.content_widget)
        layout_principal.addWidget(QLabel("Selecione os erros encontrados:"))
        layout_principal.addWidget(self.scroll, 1)

        btn_gerar = QPushButton("📄 GERAR RELATÓRIO WORD")
        btn_gerar.setStyleSheet("""
            QPushButton {
                background-color: #2196F3;
                color: white;
                font-weight: bold;
                padding: 15px;
                font-size: 16px;
                border-radius: 8px;
                margin-top: 10px;
            }
            QPushButton:hover {
                background-color: #1976D2;
            }
        """)
        btn_gerar.clicked.connect(self.gerar_relatorio)
        layout_principal.addWidget(btn_gerar)

        self.atualizar_contador()

    def carregar_equipes(self):
        equipes = listar_equipes_disponiveis()
        self.equipe.clear()
        self.equipe.addItems(equipes)

    def abrir_calendario(self):
        dialog = QDialog(self)
        dialog.setWindowTitle("Selecionar Data")
        dialog.setModal(True)
        dialog.resize(400, 300)

        layout = QVBoxLayout(dialog)

        calendario = QCalendarWidget()
        calendario.setGridVisible(True)
        calendario.setSelectedDate(QDate.currentDate())

        button_box = QDialogButtonBox(QDialogButtonBox.Ok | QDialogButtonBox.Cancel)
        button_box.accepted.connect(dialog.accept)
        button_box.rejected.connect(dialog.reject)

        layout.addWidget(calendario)
        layout.addWidget(button_box)

        if dialog.exec_() == QDialog.Accepted:
            data_selecionada = calendario.selectedDate()
            self.data.setText(data_selecionada.toString("dd/MM/yyyy"))

    def carregar_erros(self):
        df = listar_erros()
        categorias = {}

        for _, row in df.iterrows():
            categorias.setdefault(row["categoria"], []).append(row["descricao"])

        for categoria in sorted(categorias.keys()):
            self.combo_filtro.addItem(categoria)

        for categoria in sorted(categorias.keys()):
            grupo = self.criar_grupo_categoria(categoria, categorias[categoria])
            self.content_layout.addWidget(grupo)

        self.content_layout.addStretch()

    def criar_grupo_categoria(self, categoria, erros):
        grupo = QGroupBox()
        grupo.setStyleSheet("""
            QGroupBox {
                font-weight: bold;
                font-size: 14px;
                border: 2px solid #BDBDBD;
                border-radius: 8px;
                margin-top: 10px;
                padding-top: 10px;
            }
        """)

        layout_grupo = QVBoxLayout(grupo)

        cabecalho = QFrame()
        cabecalho.setStyleSheet("""
            QFrame {
                background-color: #EEEEEE;
                border-radius: 5px;
                padding: 5px;
            }
        """)
        layout_cabecalho = QHBoxLayout(cabecalho)

        btn_toggle = QPushButton("▼")
        btn_toggle.setFixedSize(30, 30)
        btn_toggle.setStyleSheet("""
            QPushButton {
                background-color: #2196F3;
                color: white;
                font-weight: bold;
                border-radius: 3px;
            }
            QPushButton:hover {
                background-color: #1976D2;
            }
        """)

        lbl_titulo = QLabel(f"{categoria} ({len(erros)} erros)")
        lbl_titulo.setStyleSheet("font-weight: bold; font-size: 13px; color: #424242;")

        checkbox_todos = QCheckBox("Selecionar todos")
        checkbox_todos.setStyleSheet("font-size: 12px; color: #616161;")

        layout_cabecalho.addWidget(btn_toggle)
        layout_cabecalho.addWidget(lbl_titulo)
        layout_cabecalho.addStretch()
        layout_cabecalho.addWidget(checkbox_todos)

        layout_grupo.addWidget(cabecalho)

        container_erros = QWidget()
        container_erros.setVisible(True)
        layout_checkboxes = QVBoxLayout(container_erros)
        layout_checkboxes.setSpacing(2)

        for erro in sorted(erros):
            checkbox = ErroCheckBox(erro, categoria)
            checkbox.toggled.connect(self.on_erro_selecionado)
            self.checkboxes.append(checkbox)

            checkbox_todos.stateChanged.connect(
                lambda state, cb=checkbox: cb.setChecked(state == Qt.Checked)
            )

            layout_checkboxes.addWidget(checkbox)

        layout_grupo.addWidget(container_erros)

        grupo.container_erros = container_erros
        grupo.btn_toggle = btn_toggle

        btn_toggle.clicked.connect(
            lambda checked, cont=container_erros, btn=btn_toggle: self.toggle_categoria(cont, btn)
        )
        grupo.categoria = categoria
        return grupo
    def toggle_categoria(self, container, botao):
        expandido = container.isVisible()
        container.setVisible(not expandido)
        botao.setText("▲" if expandido else "▼")

    def expandir_tudo(self):
        for i in range(self.content_layout.count()):
            widget = self.content_layout.itemAt(i).widget()
            if hasattr(widget, 'container_erros'):
                widget.container_erros.setVisible(True)
                if hasattr(widget, 'btn_toggle'):
                    widget.btn_toggle.setText("▲")

    def recolher_tudo(self):
        for i in range(self.content_layout.count()):
            widget = self.content_layout.itemAt(i).widget()
            if hasattr(widget, 'container_erros'):
                widget.container_erros.setVisible(False)
                if hasattr(widget, 'btn_toggle'):
                    widget.btn_toggle.setText("▼")

    def filtrar_erros(self, filtro):
        for i in range(self.content_layout.count()):
            widget = self.content_layout.itemAt(i).widget()
            if hasattr(widget, 'categoria'):  # agora os grupos têm categoria
                if filtro == "Todas as categorias" or widget.categoria == filtro:
                    widget.setVisible(True)
                else:
                    widget.setVisible(False)

    def on_erro_selecionado(self, checked):
        checkbox = self.sender()
        if checked:
            self.erros_selecionados.add(checkbox.descricao)
        else:
            self.erros_selecionados.discard(checkbox.descricao)
        self.atualizar_contador()

    def atualizar_contador(self):
        quantidade = len(self.erros_selecionados)
        self.lbl_contador.setText(f"✅ {quantidade} erro(s) selecionado(s)")
        
        if quantidade == 0:
            self.contador_frame.setStyleSheet("""
                QFrame {
                    background-color: #FFF3E0;
                    border: 2px solid #FF9800;
                    border-radius: 5px;
                    padding: 10px;
                }
            """)
        else:
            self.contador_frame.setStyleSheet("""
                QFrame {
                    background-color: #E8F5E9;
                    border: 2px solid #4CAF50;
                    border-radius: 5px;
                    padding: 10px;
                }
            """)

    def limpar_selecao(self):
        for checkbox in self.checkboxes:
            checkbox.setChecked(False)
        self.erros_selecionados.clear()
        self.atualizar_contador()

    def abrir_configuracao(self):
        dialog = ConfiguracaoDialog(self)
        dialog.exec_()

    def criar_documento_com_timbrado(self, equipe, data, erros, imagens_paths):
        """Cria documento Word usando o modelo de papel timbrado"""
        try:
            config = carregar_configuracao()
            modelo_path = config.get("modelo_docx", "")
            
            if not modelo_path or not os.path.exists(modelo_path):
                print("⚠️ Modelo não configurado ou não encontrado. Usando formato básico.")
                return self.criar_documento_basico(equipe, data, erros, imagens_paths)
            
            print(f"📄 Usando modelo de papel timbrado: {modelo_path}")
            
            # Carregar o modelo
            doc = Document(modelo_path)
            
            # Título do relatório
            titulo = doc.add_heading('RELATÓRIO DE GRAVAÇÕES', 0)
            titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Adicionar informações básicas
            doc.add_paragraph()
            p_info = doc.add_paragraph()
            p_info.add_run(f"Equipe: {equipe}").bold = True
            p_info.add_run("\n")
            p_info.add_run(f"Data: {data}").bold = True
            
            doc.add_paragraph()
            
            # Seção de ERROS/OCORRÊNCIAS
            if erros:
                titulo_erros = doc.add_paragraph()
                titulo_erros.add_run('ERROS/OCORRÊNCIAS IDENTIFICADAS:').bold = True
                
                for n, erro in enumerate(sorted(erros), 1):
                    p_erro = doc.add_paragraph()
                    p_erro.add_run(f"{n}. {erro}")
            else:
                p_sem_erros = doc.add_paragraph()
                p_sem_erros.add_run("✅ Nenhuma ocorrência registrada").italic = True
            
            doc.add_paragraph()
            
            # Seção de EVIDÊNCIAS FOTOGRÁFICAS
            titulo_fotos = doc.add_paragraph()
            titulo_fotos.add_run('EVIDÊNCIAS FOTOGRÁFICAS:').bold = True
            
            if imagens_paths and len(imagens_paths) > 0:
                print(f"📷 Incluindo {len(imagens_paths)} imagens no relatório...")
                
                for img_path in imagens_paths:
                    if os.path.exists(img_path):
                        try:
                            # Adicionar nome da imagem
                            p_nome_img = doc.add_paragraph()
                            p_nome_img.add_run(f"• {os.path.basename(img_path)}").italic = True
                            
                            # Adicionar imagem
                            doc.add_picture(img_path, width=Inches(5.5))  # 5.5 polegadas
                            
                            # Adicionar espaço após imagem
                            doc.add_paragraph()
                            
                        except Exception as e:
                            print(f"⚠️ Erro ao adicionar imagem {img_path}: {e}")
                            p_erro = doc.add_paragraph()
                            p_erro.add_run(f"[Imagem não pode ser carregada: {os.path.basename(img_path)}]").italic = True
                    else:
                        p_erro = doc.add_paragraph()
                        p_erro.add_run(f"[Arquivo não encontrado: {img_path}]").italic = True
            else:
                p_sem_fotos = doc.add_paragraph()
                p_sem_fotos.add_run('*--- Nenhuma evidência fotográfica anexada').italic = True
            
            # Linha separadora
            doc.add_paragraph('――――――――――――――――――――――――――――――――――――――――――――――――――――――――――――')
            
            # Rodapé
            rodape = doc.add_paragraph()
            
            data_geracao = datetime.now().strftime('%d/%m/%Y %H:%M')
            rodape_texto = f"📋 Relatório gerado em {data_geracao} | 👥 {equipe} | 📅 {data} | ⚡ Sistema ENGELMIG"
            rodape_run = rodape.add_run(rodape_texto)
            rodape.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            rodape_run.font.size = Pt(9)
            rodape_run.font.color.rgb = RGBColor(100, 100, 100)
            
            return doc
            
        except Exception as e:
            print(f"❌ Erro ao criar documento com timbrado: {e}")
            import traceback
            traceback.print_exc()
            # Fallback para documento básico
            return self.criar_documento_basico(equipe, data, erros, imagens_paths)

    def criar_documento_basico(self, equipe, data, erros, imagens_paths):
        """Cria documento Word básico (fallback)"""
        try:
            doc = Document()
            
            # Configurar margens
            for section in doc.sections:
                section.top_margin = Cm(2.54)
                section.bottom_margin = Cm(2.54)
                section.left_margin = Cm(2.54)
                section.right_margin = Cm(2.54)
            
            # Título
            titulo = doc.add_heading('RELATÓRIO DE GRAVAÇÕES', 0)
            titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Adicionar espaço
            doc.add_paragraph()
            
            # Equipe
            p_equipe = doc.add_paragraph()
            p_equipe.add_run('Equipe: ').bold = True
            p_equipe.add_run(equipe)
            
            # Data
            p_data = doc.add_paragraph()
            p_data.add_run('Data do Relatório: ').bold = True
            p_data.add_run(data)
            
            # Adicionar espaço
            doc.add_paragraph()
            
            # ERROS/OCORRÊNCIAS
            titulo_erros = doc.add_paragraph()
            titulo_erros.add_run('ERROS/OCORRÊNCIAS:').bold = True
            
            # Adicionar espaço
            doc.add_paragraph()
            
            if erros:
                for n, erro in enumerate(sorted(erros), 1):
                    p_erro = doc.add_paragraph()
                    p_erro.add_run(f"{n}. {erro}")
            else:
                p_nenhum = doc.add_paragraph("Nenhuma ocorrência registrada")
                p_nenhum.runs[0].italic = True
            
            # Adicionar espaço
            doc.add_paragraph()
            
            # EVIDÊNCIAS FOTOGRÁFICAS
            titulo_fotos = doc.add_paragraph()
            titulo_fotos.add_run('EVIDÊNCIAS FOTOGRÁFICAS:').bold = True
            
            # Adicionar espaço
            doc.add_paragraph()
            
            if imagens_paths and len(imagens_paths) > 0:
                print(f"Incluindo {len(imagens_paths)} imagens no relatório...")
                for img_path in imagens_paths:
                    if os.path.exists(img_path):
                        try:
                            # Adicionar nome do arquivo
                            p_nome = doc.add_paragraph(f"• {os.path.basename(img_path)}")
                            
                            # Adicionar imagem
                            doc.add_picture(img_path, width=Inches(6))
                            
                            # Adicionar espaço entre imagens
                            doc.add_paragraph()
                            
                        except Exception as e:
                            print(f"Erro ao adicionar imagem {img_path}: {e}")
                            doc.add_paragraph(f"[Erro ao carregar imagem: {os.path.basename(img_path)}]")
                    else:
                        doc.add_paragraph(f"[Arquivo não encontrado: {img_path}]")
            else:
                p_sem_fotos = doc.add_paragraph()
                p_sem_fotos.add_run('*--- Nenhuma evidência fotográfica anexada').italic = True
            
            # Adicionar linha separadora
            doc.add_paragraph('――――――――――――――――――――――――――――――――――――――――――――――――――――――――――――')
            
            # Rodapé informativo
            rodape = doc.add_paragraph()
            
            data_geracao = datetime.now().strftime('%d/%m/%Y %H:%M')
            rodape_texto = f"📋 Relatório gerado em {data_geracao} | 👥 {equipe} | 📅 {data} | ⚡ Sistema ENGELMIG"
            rodape_run = rodape.add_run(rodape_texto)
            rodape.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Configurar fonte do rodapé
            rodape_run.font.size = Pt(9)
            rodape_run.font.color.rgb = RGBColor(100, 100, 100)  # Cinza
            
            return doc
            
        except Exception as e:
            print(f"Erro ao criar documento: {e}")
            import traceback
            traceback.print_exc()
            raise

    def gerar_relatorio(self):
        equipe = self.equipe.currentText().strip()
        data = self.data.text().strip()
        status = "Conforme" if self.btn_conforme.isChecked() else "Não Conforme"
        erros_lista = list(self.erros_selecionados)

        if not equipe:
            QMessageBox.warning(self, "Atenção", "Por favor, informe o nome da equipe.")
            self.equipe.setFocus()
            return
            
        if not data:
            QMessageBox.warning(self, "Atenção", "Por favor, informe a data.")
            self.data.setFocus()
            return

        try:
            datetime.strptime(data, "%d/%m/%Y")
        except ValueError:
            QMessageBox.warning(self, "Data Inválida",
                            "Por favor, informe a data no formato DD/MM/AAAA.")
            self.data.setFocus()
            return

        if status == "Conforme" and self.erros_selecionados:
            resposta = QMessageBox.question(
                self, "Confirmar Status",
                "Você selecionou erros mas marcou como 'Conforme'.\n\n"
                "Deseja continuar assim ou mudar para 'Não Conforme'?",
                QMessageBox.Yes | QMessageBox.No | QMessageBox.Cancel,
                QMessageBox.No
            )
            
            if resposta == QMessageBox.Cancel:
                return
            elif resposta == QMessageBox.No:
                status = "Não Conforme"
                self.btn_nao_conforme.setChecked(True)

        if not self.erros_selecionados:
            resposta = QMessageBox.question(
                self, "Confirmar",
                f"Nenhum erro foi selecionado. Status: {status}\n\n"
                "Deseja gerar o relatório mesmo assim?",
                QMessageBox.Yes | QMessageBox.No,
            )
            if resposta == QMessageBox.No:
                return

        erros_lista = sorted(list(self.erros_selecionados))
        
        try:
            relatorio_id = salvar_relatorio(equipe, data, erros_lista, status)
            print(f"Relatório salvo no banco com ID: {relatorio_id}")
        except Exception as e:
            QMessageBox.critical(self, "Erro", f"Erro ao salvar no banco de dados:\n{str(e)}")
            return

        try:
            # Obter imagens do App principal
            imagens_para_relatorio = []
            if hasattr(self.parent_app, 'lista_imagens_compartilhada'):
                imagens_para_relatorio = self.parent_app.lista_imagens_compartilhada.copy()
                print(f"=== IMAGENS PARA O RELATÓRIO ===")
                print(f"Número de imagens: {len(imagens_para_relatorio)}")
                for i, img in enumerate(imagens_para_relatorio, 1):
                    print(f"{i}. {img} (existe: {os.path.exists(img)})")
                print("================================")
            
            # Criar documento usando papel timbrado
            doc = self.criar_documento_com_timbrado(equipe, data, erros_lista, imagens_para_relatorio)

            # Sugerir nome do arquivo
            nome_sugerido = f"Relatório_{equipe}_{data.replace('/', '-')}.docx"
            
            # Pedir para salvar
            path, _ = QFileDialog.getSaveFileName(
                self, "Salvar Relatório",
                nome_sugerido,
                "Documentos Word (*.docx)"
            )

            if path:
                # Salvar documento
                doc.save(path)
                print(f"✅ Documento salvo em: {path}")
            self.limpar_dados_apos_relatorio()

        except Exception as e:
            QMessageBox.critical(
                self,
                "Erro ao gerar relatório",
                f"Ocorreu um erro ao gerar o relatório:\n\n{str(e)}\n\n"
                f"Detalhes técnicos: {type(e).__name__}"
            )
            import traceback
            traceback.print_exc()
        
    
    def limpar_dados_apos_relatorio(self):
        """Limpa todos os dados após gerar um relatório"""
        try:
            # 1. Limpar seleção de erros
            for checkbox in self.checkboxes:
                checkbox.setChecked(False)
            self.erros_selecionados.clear()
            
            # 2. Limpar lista de imagens no App principal
            if hasattr(self.parent_app, 'lista_imagens_compartilhada'):
                self.parent_app.lista_imagens_compartilhada.clear()
                print("Lista de imagens limpa após gerar relatório")
            
             # 3. Atualizar contador
            self.atualizar_contador()
            
            print("✅ Dados limpos após gerar relatório")
            
        except Exception as e:
            print(f"Erro ao limpar dados: {e}")

class AbaImagens(QWidget):
    def __init__(self, parent_app):
        super().__init__()
        self.parent_app = parent_app
        self.temp_dir = os.path.join(BASE_DIR, "temp_images")
        os.makedirs(self.temp_dir, exist_ok=True)
        self.init_ui()
        
        print(f"AbaImagens inicializada com suporte a colagem (Ctrl+V)")
    
    def init_ui(self):
        layout = QVBoxLayout(self)

        cabecalho = QLabel("🖼️ EVIDÊNCIAS FOTOGRÁFICAS")
        cabecalho.setStyleSheet("""
            QLabel {
                font-size: 18px;
                font-weight: bold;
                color: #388E3C;
                padding: 10px;
                background-color: #E8F5E9;
                border-radius: 5px;
            }
        """)
        cabecalho.setAlignment(Qt.AlignCenter)
        layout.addWidget(cabecalho)

        # Dica de colagem
        lbl_dica = QLabel("💡 Dica: Cole imagens diretamente aqui (Ctrl+V)")
        lbl_dica.setStyleSheet("color: #2196F3; font-style: italic; padding: 5px;")
        lbl_dica.setAlignment(Qt.AlignCenter)
        layout.addWidget(lbl_dica)

        painel_controle = QFrame()
        painel_controle.setStyleSheet("""
            QFrame {
                background-color: #F5F5F5;
                border-radius: 5px;
                padding: 10px;
            }
        """)

        layout_controle = QHBoxLayout(painel_controle)

        btn_add = QPushButton("📁 Adicionar Imagens")
        btn_add.setStyleSheet("""
            QPushButton {
                background-color: #4CAF50;
                color: white;
                padding: 10px;
                font-weight: bold;
                border-radius: 5px;
            }
            QPushButton:hover {
                background-color: #45a049;
            }
        """)
        btn_add.clicked.connect(self.adicionar_imagens)

        btn_rem = QPushButton("🗑️ Remover Selecionada")
        btn_rem.setStyleSheet("""
            QPushButton {
                background-color: #f44336;
                color: white;
                padding: 10px;
                font-weight: bold;
                border-radius: 5px;
            }
            QPushButton:hover {
                background-color: #d32f2f;
            }
        """)
        btn_rem.clicked.connect(self.remover_selecionada)

        btn_limpar = QPushButton("🧹 Limpar Todas")
        btn_limpar.setStyleSheet("""
            QPushButton {
                background-color: #FF9800;
                color: white;
                padding: 10px;
                font-weight: bold;
                border-radius: 5px;
            }
            QPushButton:hover {
                background-color: #F57C00;
            }
        """)
        btn_limpar.clicked.connect(self.limpar_todas)

        layout_controle.addWidget(btn_add)
        layout_controle.addWidget(btn_rem)
        layout_controle.addWidget(btn_limpar)
        layout_controle.addStretch()

        layout.addWidget(painel_controle)

        self.lbl_contador = QLabel("🖼️ 0 imagem(s) selecionada(s)")
        self.lbl_contador.setStyleSheet("""
            QLabel {
                font-weight: bold;
                font-size: 14px;
                color: #2196F3;
                padding: 5px;
            }
        """)
        layout.addWidget(self.lbl_contador)

        self.lista = QListWidget()
        self.lista.setIconSize(QSize(150, 150))
        self.lista.setResizeMode(QListWidget.Adjust)
        self.lista.setSpacing(5)
        self.lista.setStyleSheet("""
            QListWidget {
                background-color: white;
                border: 1px solid #E0E0E0;
                border-radius: 5px;
            }
            QListWidget::item {
                padding: 10px;
                border-bottom: 1px solid #EEEEEE;
            }
            QListWidget::item:selected {
                background-color: #E3F2FD;
            }
        """)
        layout.addWidget(self.lista)

        self.atualizar_contador()
        self.atualizar_lista()
    
    def obter_imagens(self):
        """Obtém a lista de imagens do App principal"""
        if hasattr(self.parent_app, 'lista_imagens_compartilhada'):
            return self.parent_app.lista_imagens_compartilhada
        return []
    
    def adicionar_imagens(self):
        config = carregar_configuracao()
        ultimo_dir = config.get("ultimo_diretorio", BASE_DIR)
        
        files, _ = QFileDialog.getOpenFileNames(
            self, "Selecionar Imagens", ultimo_dir,
            "Imagens (*.png *.jpg *.jpeg *.bmp *.gif);;Todos os arquivos (*.*)"
        )
        
        if files:
            config["ultimo_diretorio"] = os.path.dirname(files[0])
            salvar_configuracao(config)
        
        novas_imagens = 0
        for f in files:
            if os.path.exists(f) and f not in self.obter_imagens():
                self.obter_imagens().append(f)
                novas_imagens += 1
                print(f"Imagem adicionada: {f}")
        
        if novas_imagens > 0:
            self.atualizar_lista()
            print(f"Total de imagens após adição: {len(self.obter_imagens())}")
    
    def colar_imagem(self):
        """Cola uma imagem da área de transferência e salva como arquivo"""
        clipboard = QApplication.clipboard()
        mime_data = clipboard.mimeData()
        
        if mime_data.hasImage():
            # Obtém a imagem
            qimage = clipboard.image()
            if qimage.isNull():
                QMessageBox.warning(self, "Erro", "Nenhuma imagem válida na área de transferência.")
                return
            
            # Gera nome único
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            nome_arquivo = f"colado_{timestamp}.png"
            caminho_completo = os.path.join(self.temp_dir, nome_arquivo)
            
            # Salva a imagem
            if qimage.save(caminho_completo, "PNG"):
                # Adiciona à lista compartilhada
                if caminho_completo not in self.obter_imagens():
                    self.obter_imagens().append(caminho_completo)
                    self.atualizar_lista()
                else:
                    QMessageBox.information(self, "Aviso", "Imagem já está na lista.")
            else:
                QMessageBox.critical(self, "Erro", "Não foi possível salvar a imagem colada.")
        else:
            QMessageBox.warning(self, "Sem imagem", "A área de transferência não contém uma imagem.\n\nDica: Tire um print (PrintScreen) ou copie uma imagem de outro lugar antes de colar.")
    
    def keyPressEvent(self, event):
        """Captura Ctrl+V para colar imagens"""
        if event.key() == Qt.Key_V and event.modifiers() == Qt.ControlModifier:
            self.colar_imagem()
        else:
            super().keyPressEvent(event)
    
    def remover_selecionada(self):
        item = self.lista.currentItem()
        if item:
            caminho = item.data(Qt.UserRole)
            if caminho in self.obter_imagens():
                self.obter_imagens().remove(caminho)
                self.atualizar_lista()
                print(f"Imagem removida: {caminho}")
    
    def limpar_todas(self):
        if len(self.obter_imagens()) > 0:
            resposta = QMessageBox.question(
                self, "Confirmar",
                f"Remover todas as {len(self.obter_imagens())} imagens?",
                QMessageBox.Yes | QMessageBox.No
            )
            if resposta == QMessageBox.Yes:
                self.obter_imagens().clear()
                self.atualizar_lista()
                print("Todas as imagens foram removidas")
    
    def atualizar_lista(self):
        self.lista.clear()
        imagens = self.obter_imagens()

        for img_path in imagens:
            if os.path.exists(img_path):
                nome = os.path.basename(img_path)
                item = QListWidgetItem(nome)
                item.setToolTip(img_path)

                try:
                    pix = QPixmap(img_path)
                    if not pix.isNull():
                        pix = pix.scaled(150, 150, Qt.KeepAspectRatio, Qt.SmoothTransformation)
                        item.setIcon(QIcon(pix))
                except:
                    item.setIcon(QIcon())

                item.setData(Qt.UserRole, img_path)
                self.lista.addItem(item)
            else:
                # Se o arquivo não existe mais, remova da lista
                self.obter_imagens().remove(img_path)
        
        self.atualizar_contador()
    
    def atualizar_contador(self):
        total = len(self.obter_imagens())
        self.lbl_contador.setText(f"🖼️ {total} imagem(s) selecionada(s)")

class AbaRegistro(QWidget):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.parent_app = parent
    
    def init_ui(self):
        layout = QVBoxLayout(self)

        cabecalho = QLabel("📋 REGISTRO DE RELATÓRIOS")
        cabecalho.setStyleSheet("""
            QLabel {
                font-size: 18px;
                font-weight: bold;
                color: #7B1FA2;
                padding: 15px;
                background-color: #F3E5F5;
                border-radius: 5px;
            }
        """)
        cabecalho.setAlignment(Qt.AlignCenter)
        layout.addWidget(cabecalho)

        painel_filtros = QFrame()
        painel_filtros.setStyleSheet("""
            QFrame {
                background-color: #F5F5F5;
                border-radius: 5px;
                padding: 10px;
            }
        """)

        filtros_layout = QGridLayout(painel_filtros)
        filtros_layout.setSpacing(10)
        lbl_equipe = QLabel("Filtrar por Equipe:")
        self.filtro_equipe = QComboBox()
        self.filtro_equipe.setEditable(True)
        self.filtro_equipe.addItem("Todas as equipes")
        self.filtro_equipe.setStyleSheet("""
            QComboBox {
                padding: 5px;
                border: 1px solid #BDBDBD;
                border-radius: 3px;
            }
        """)

        lbl_data = QLabel("Filtrar por Data:")
        self.filtro_data = QComboBox()
        self.filtro_data.addItem("Todas as datas")
        self.filtro_data.setStyleSheet("""
            QComboBox {
                padding: 5px;
                border: 1px solid #BDBDBD;
                border-radius: 3px;
            }
        """)

        lbl_status = QLabel("Filtrar por Status:")
        self.filtro_status = QComboBox()
        self.filtro_status.addItem("Todos os status")
        self.filtro_status.setStyleSheet("""
            QComboBox {
                padding: 5px;
                border: 1px solid #BDBDBD;
                border-radius: 3px;
            }
        """)

        lbl_erro = QLabel("Filtrar por Erro:")
        self.filtro_erro = QLineEdit()
        self.filtro_erro.setPlaceholderText("Digite parte do erro...")
        self.filtro_erro.setStyleSheet("""
            QLineEdit {
                padding: 5px;
                border: 1px solid #BDBDBD;
                border-radius: 3px;
            }
        """)

        btn_exportar = QPushButton("Exportar para Excel")
        btn_exportar.setStyleSheet("""
            QPushButton {
                background-color: #4CAF50;
                color: white;
                padding: 8px 15px;
                border-radius: 3px;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #45a049;
            }
        """)
        btn_exportar.clicked.connect(self.exportar_para_excel)
        filtros_layout.addWidget(btn_exportar, 2, 4)

        btn_buscar = QPushButton("Buscar")
        btn_buscar.setStyleSheet("""
            QPushButton {
                background-color: #2196F3;
                color: white;
                padding: 8px 15px;
                border-radius: 3px;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #1976D2;
            }
        """)
        btn_buscar.clicked.connect(self.aplicar_filtros)

        btn_limpar = QPushButton("🧹 Limpar Filtros")
        btn_limpar.setStyleSheet("""
            QPushButton {
                background-color: #757575;
                color: white;
                padding: 8px 15px;
                border-radius: 3px;
            }
            QPushButton:hover {
                background-color: #616161;
            }
        """)
        btn_limpar.clicked.connect(self.limpar_filtros)

        filtros_layout.addWidget(lbl_equipe, 0, 0)
        filtros_layout.addWidget(self.filtro_equipe, 0, 1)
        filtros_layout.addWidget(lbl_data, 0, 2)
        filtros_layout.addWidget(self.filtro_data, 0, 3)
        filtros_layout.addWidget(lbl_status, 1, 0)
        filtros_layout.addWidget(self.filtro_status, 1, 1)
        filtros_layout.addWidget(lbl_erro, 1, 2)
        filtros_layout.addWidget(self.filtro_erro, 1, 3)
        filtros_layout.addWidget(btn_buscar, 2, 2)
        filtros_layout.addWidget(btn_limpar, 2, 3)

        layout.addWidget(painel_filtros)

        self.tabela = QTableWidget()
        self.tabela.setColumnCount(7)
        self.tabela.setHorizontalHeaderLabels([
            "ID", "Equipe", "Data", "Status", "Data Criação", "Total Erros", "Ações"
        ])

        # Configurar redimensionamento das colunas
        header = self.tabela.horizontalHeader()
        header.setSectionResizeMode(0, QHeaderView.ResizeToContents)  # ID
        header.setSectionResizeMode(1, QHeaderView.Stretch)           # Equipe (expande)
        header.setSectionResizeMode(2, QHeaderView.ResizeToContents)  # Data
        header.setSectionResizeMode(3, QHeaderView.ResizeToContents)  # Status
        header.setSectionResizeMode(4, QHeaderView.ResizeToContents)  # Data Criação
        header.setSectionResizeMode(5, QHeaderView.ResizeToContents)  # Total Erros
        header.setSectionResizeMode(6, QHeaderView.Fixed)             # Ações (tamanho fixo)
        self.tabela.setColumnWidth(6, 280)  # Largura suficiente para os 4 botões

        self.tabela.setSelectionBehavior(QAbstractItemView.SelectRows)
        self.tabela.setAlternatingRowColors(True)
        self.tabela.setStyleSheet("""
            QTableWidget {
                background-color: white;
                alternate-background-color: #F5F5F5;
                gridline-color: #E0E0E0;
            }
            QTableWidget::item {
                padding: 10px;
            }
            QTableWidget::item:selected {
                background-color: #E3F2FD;
            }
            QHeaderView::section {
                background-color: #1976D2;
                color: white;
                font-weight: bold;
                padding: 15px;
                border: none;
            }
        """)

        layout.addWidget(self.tabela, 1)

        self.lbl_info = QLabel("Total de registros: 0")
        self.lbl_info.setStyleSheet("""
            QLabel {
                font-weight: bold;
                color: #424242;
                padding: 10px;
            }
        """)
        layout.addWidget(self.lbl_info)

        self.carregar_dados_filtros()
        self.carregar_registros()
        self.tabela.cellDoubleClicked.connect(self.on_cell_double_clicked)

    def carregar_dados_filtros(self):
        equipes = listar_equipes_disponiveis()
        self.filtro_equipe.clear()
        self.filtro_equipe.addItem("Todas as equipes")
        self.filtro_equipe.addItems(equipes)

        datas = listar_datas_disponiveis()
        self.filtro_data.clear()
        self.filtro_data.addItem("Todas as datas")
        self.filtro_data.addItems(datas)
        
        status_lista = listar_status_disponiveis()
        self.filtro_status.clear()
        self.filtro_status.addItem("Todos os status")
        self.filtro_status.addItems(status_lista)
    
    def aplicar_filtros(self):
        self.carregar_registros()
    
    def limpar_filtros(self):
        self.filtro_equipe.setCurrentIndex(0)
        self.filtro_data.setCurrentIndex(0)
        self.filtro_status.setCurrentIndex(0)
        self.filtro_erro.clear()
        self.carregar_registros()
    
    def carregar_registros(self):
        import pandas as pd
        
        filtro_equipe = "" if self.filtro_equipe.currentText() == "Todas as equipes" else self.filtro_equipe.currentText()
        filtro_data = "" if self.filtro_data.currentText() == "Todas as datas" else self.filtro_data.currentText()
        filtro_status = "" if self.filtro_status.currentText() == "Todos os status" else self.filtro_status.currentText()
        filtro_erro = self.filtro_erro.text().strip()

        df = listar_relatorios(filtro_equipe, filtro_data, filtro_erro, filtro_status)

        self.tabela.setRowCount(len(df))

        for i, row in df.iterrows():
            item_id = QTableWidgetItem(str(row['id']))
            item_id.setTextAlignment(Qt.AlignCenter)
            self.tabela.setItem(i, 0, item_id)

            item_equipe = QTableWidgetItem(row['equipe'])
            item_equipe.setToolTip("Clique duas vezes para editar")
            self.tabela.setItem(i, 1, item_equipe)
            
            self.tabela.setItem(i, 2, QTableWidgetItem(row['data']))
            
            item_status = QTableWidgetItem(row['status'])
            if row['status'] == 'Conforme':
                item_status.setForeground(Qt.darkGreen)
            else:
                item_status.setForeground(Qt.darkRed)
            item_status.setTextAlignment(Qt.AlignCenter)
            self.tabela.setItem(i, 3, item_status)
            
            if 'data_criacao' in row and pd.notna(row['data_criacao']):
                data_criacao = str(row['data_criacao'])
                if ' ' in data_criacao:
                    data_criacao = data_criacao.split(' ')[0]
            else:
                data_criacao = "N/A"
            self.tabela.setItem(i, 4, QTableWidgetItem(data_criacao))
            
            item_total = QTableWidgetItem(str(row['total_erros']))
            item_total.setTextAlignment(Qt.AlignCenter)
            self.tabela.setItem(i, 5, item_total)

            # Ajustar altura da linha para caber os botões
            self.tabela.setRowHeight(i, 45)

            widget_acoes = QWidget()
            layout_acoes = QHBoxLayout(widget_acoes)
            layout_acoes.setContentsMargins(2, 2, 2, 2)
            layout_acoes.setSpacing(2)

            btn_ver = QPushButton("👁️ Ver")
            btn_ver.setFixedSize(60, 30)
            btn_ver.setStyleSheet("""
                QPushButton {
                    background-color: #2196F3;
                    color: white;
                    border-radius: 3px;
                    font-size: 11px;
                    padding: 5px 8px;
                }
                QPushButton:hover {
                    background-color: #1976D2;
                }
            """)
            btn_ver.clicked.connect(lambda checked, rid=row['id']: self.ver_detalhes(rid))

            btn_editar_equipe = QPushButton("✏️ Equipe")
            btn_editar_equipe.setFixedSize(70, 30)
            btn_editar_equipe.setToolTip("Editar nome da equipe")
            btn_editar_equipe.setStyleSheet("""
                QPushButton {
                    background-color: #9C27B0;
                    color: white;
                    border-radius: 3px;
                    font-size: 11px;
                    padding: 5px 8px;
                }
                QPushButton:hover {
                    background-color: #7B1FA2;
                }
            """)
            btn_editar_equipe.clicked.connect(lambda checked, rid=row['id'], equipe=row['equipe']: self.editar_equipe(rid, equipe))

            btn_editar_status = QPushButton("✏️ Status")
            btn_editar_status.setFixedSize(70, 30)
            btn_editar_status.setToolTip("Alterar status")
            btn_editar_status.setStyleSheet("""
                QPushButton {
                    background-color: #FF9800;
                    color: white;
                    border-radius: 3px;
                    font-size: 11px;
                    padding: 5px 8px;
                }
                QPushButton:hover {
                    background-color: #F57C00;
                }
            """)
            btn_editar_status.clicked.connect(lambda checked, rid=row['id'], status=row['status']: self.editar_status(rid, status))

            btn_excluir = QPushButton("🗑️")
            btn_excluir.setFixedSize(30, 30)
            btn_excluir.setToolTip("Excluir relatório")
            btn_excluir.setStyleSheet("""
                QPushButton {
                    background-color: #f44336;
                    color: white;
                    border-radius: 3px;
                    font-size: 12px;
                    padding: 5px;
                }
                QPushButton:hover {
                    background-color: #d32f2f;
                }
            """)
            btn_excluir.clicked.connect(lambda checked, rid=row['id']: self.excluir_registro(rid))

            layout_acoes.addWidget(btn_ver)
            layout_acoes.addWidget(btn_editar_equipe)
            layout_acoes.addWidget(btn_editar_status)
            layout_acoes.addWidget(btn_excluir)
            layout_acoes.addStretch()

            self.tabela.setCellWidget(i, 6, widget_acoes)

        self.lbl_info.setText(f"Total de registros: {len(df)}")
    
    def editar_status(self, relatorio_id, status_atual):
        dialog = QDialog(self)
        dialog.setWindowTitle(f"Alterar Status - Relatório #{relatorio_id}")
        dialog.setModal(True)
        dialog.resize(400, 100)
        
        layout = QVBoxLayout(dialog)
        
        lbl_titulo = QLabel(f"Alterar status do relatório #{relatorio_id}")
        lbl_titulo.setStyleSheet("font-weight: bold; font-size: 14px;")
        lbl_titulo.setAlignment(Qt.AlignCenter)
        layout.addWidget(lbl_titulo)
        
        lbl_status_atual = QLabel(f"Status atual: {status_atual}")
        if status_atual == "Conforme":
            lbl_status_atual.setStyleSheet("color: #388E3C; font-weight: bold;")
        else:
            lbl_status_atual.setStyleSheet("color: #D32F2F; font-weight: bold;")
        layout.addWidget(lbl_status_atual)
        
        lbl_novo_status = QLabel("Novo status:")
        layout.addWidget(lbl_novo_status)
        
        status_group = QButtonGroup(dialog)
        rb_conforme = QRadioButton("Conforme")
        rb_nao_conforme = QRadioButton("Não Conforme")
        
        if status_atual == "Conforme":
            rb_conforme.setChecked(True)
        else:
            rb_nao_conforme.setChecked(True)
        
        status_group.addButton(rb_conforme)
        status_group.addButton(rb_nao_conforme)
        
        layout.addWidget(rb_conforme)
        layout.addWidget(rb_nao_conforme)
        
        button_box = QDialogButtonBox(QDialogButtonBox.Ok | QDialogButtonBox.Cancel)
        button_box.accepted.connect(dialog.accept)
        button_box.rejected.connect(dialog.reject)
        layout.addWidget(button_box)
        
        if dialog.exec_() == QDialog.Accepted:
            novo_status = "Conforme" if rb_conforme.isChecked() else "Não Conforme"
            
            if atualizar_status_relatorio(relatorio_id, novo_status):
                QMessageBox.information(
                    self,
                    "Status Atualizado",
                    f"Status do relatório #{relatorio_id} atualizado para: {novo_status}"
                )
                self.carregar_registros()
    
    def ver_detalhes(self, relatorio_id):
        try:
            df_relatorio, df_erros = obter_detalhes_relatorio(relatorio_id)
            
            if df_relatorio.empty:
                QMessageBox.warning(self, "Erro", f"Relatório #{relatorio_id} não encontrado.")
                return
            
            relatorio = df_relatorio.iloc[0]
            
            mensagem = f"RELATÓRIO #{relatorio_id}\n\n"
            mensagem += f"Equipe: {relatorio['equipe']}\n"
            mensagem += f"Data: {relatorio['data']}\n"
            mensagem += f"Status: {relatorio['status']}\n"
            mensagem += f"Criado em: {relatorio['data_criacao']}\n\n"
            
            if not df_erros.empty:
                mensagem += "ERROS IDENTIFICADOS:\n"
                for _, erro in df_erros.iterrows():
                    mensagem += f"  • {erro['descricao']}\n"
                mensagem += f"\nTotal: {len(df_erros)} erro(s)"
            else:
                mensagem += "Nenhum erro registrado"
            
            QMessageBox.information(self, "Detalhes do Relatório", mensagem)
            
        except Exception as e:
            QMessageBox.critical(self, "Erro", f"Erro ao obter detalhes:\n{str(e)}")
    
    def exportar_para_excel(self):
        from PyQt5.QtCore import QThread, pyqtSignal

        class ExportThread(QThread):
            finished = pyqtSignal(object)

            def __init__(self, func):
                super().__init__()
                self.func = func

            def run(self):
                try:
                    resultado = self.func()
                    self.finished.emit(resultado)
                except Exception as e:
                    self.finished.emit({'sucesso': False, 'erro': str(e)})
        btn = self.sender()
        btn.setEnabled(False)

        def on_finished(resultado):
            btn.setEnabled(True)
            if resultado and 'arquivo' in resultado:
                QMessageBox.information(self, "Sucesso", f"Planilha exportada com sucesso!\nArquivo: {resultado['arquivo']}")
            else:
                erro = resultado.get('erro', 'Erro desconhecido') if resultado else 'resultado é None'
                QMessageBox.critical(self, "Erro", f"Falha na exportação:\n{erro}")

        self.thread = ExportThread(exportar_dados_para_excel_fixo)
        self.thread.finished.connect(on_finished)
        self.thread.start()
    
    def excluir_registro(self, relatorio_id):
        resposta = QMessageBox.question(
            self, "Confirmar Exclusão",
            f"Deseja realmente excluir o relatório #{relatorio_id}?\n\n"
            "Esta ação não pode ser desfeita.",
            QMessageBox.Yes | QMessageBox.No,
            QMessageBox.No
        )

        if resposta == QMessageBox.Yes:
            sucesso = excluir_relatorio(relatorio_id)

            if sucesso:
                QMessageBox.information(self, "Sucesso", f"Relatório #{relatorio_id} excluído com sucesso.")
                self.carregar_registros()
            else:
                QMessageBox.warning(self, "Erro", f"Não foi possível excluir o relatório #{relatorio_id}.")
    
    def on_cell_double_clicked(self, row, column):
        if column == 1:
            item = self.tabela.item(row, column)
            if item:
                equipe_atual = item.text()
                
                id_item = self.tabela.item(row, 0)
                if id_item:
                    relatorio_id = int(id_item.text())
                    self.editar_equipe(relatorio_id, equipe_atual)
 
    def editar_equipe(self, relatorio_id, equipe_atual):
        dialog = QDialog(self)
        dialog.setWindowTitle(f"Editar Equipe - Relatório #{relatorio_id}")
        dialog.setModal(True)
        dialog.resize(400, 200)
        
        layout = QVBoxLayout(dialog)
        
        lbl_titulo = QLabel(f"Editar equipe do relatório #{relatorio_id}")
        lbl_titulo.setStyleSheet("font-weight: bold; font-size: 14px;")
        lbl_titulo.setAlignment(Qt.AlignCenter)
        layout.addWidget(lbl_titulo)
        
        lbl_equipe_atual = QLabel(f"Equipe atual: {equipe_atual}")
        lbl_equipe_atual.setStyleSheet("color: #1976D2; font-weight: bold;")
        layout.addWidget(lbl_equipe_atual)
        
        lbl_nova_equipe = QLabel("Novo nome da equipe:")
        layout.addWidget(lbl_nova_equipe)
        
        self.txt_nova_equipe = QComboBox()
        self.txt_nova_equipe.setEditable(True)
        self.txt_nova_equipe.addItem(equipe_atual)
        
        equipes_existentes = listar_equipes_disponiveis()
        for equipe in equipes_existentes:
            if equipe != equipe_atual and equipe not in [self.txt_nova_equipe.itemText(i) for i in range(self.txt_nova_equipe.count())]:
                self.txt_nova_equipe.addItem(equipe)
        
        self.txt_nova_equipe.setCurrentText(equipe_atual)
        self.txt_nova_equipe.setStyleSheet("""
            QComboBox {
                padding: 8px;
                font-size: 14px;
                border: 2px solid #2196F3;
                border-radius: 4px;
            }
        """)
        layout.addWidget(self.txt_nova_equipe)
        
        button_box = QDialogButtonBox(QDialogButtonBox.Ok | QDialogButtonBox.Cancel)
        button_box.accepted.connect(dialog.accept)
        button_box.rejected.connect(dialog.reject)
        layout.addWidget(button_box)
        
        if dialog.exec_() == QDialog.Accepted:
            nova_equipe = self.txt_nova_equipe.currentText().strip()
            
            if not nova_equipe:
                QMessageBox.warning(
                    self,
                    "Nome Inválido",
                    "O nome da equipe não pode estar vazio."
                )
                return
            
            if nova_equipe == equipe_atual:
                QMessageBox.information(
                    self,
                    "Nenhuma Alteração",
                    "O nome da equipe não foi alterado."
                )
                return
            
            resposta = QMessageBox.question(
                self,
                "Confirmar Alteração",
                f"Deseja alterar a equipe do relatório #{relatorio_id}?\n\n"
                f"De: {equipe_atual}\n"
                f"Para: {nova_equipe}",
                QMessageBox.Yes | QMessageBox.No,
                QMessageBox.Yes
            )
            
            if resposta == QMessageBox.Yes:
                if atualizar_nome_equipe(relatorio_id, nova_equipe):
                    QMessageBox.information(
                        self,
                        "Equipe Atualizada",
                        f"✅ Equipe do relatório #{relatorio_id} atualizada:\n\n"
                        f"Antes: {equipe_atual}\n"
                        f"Depois: {nova_equipe}"
                    )
                    self.carregar_registros()
                    self.carregar_dados_filtros()

class AutomacaoThread(QThread):
    log_signal = pyqtSignal(str)
    finished_signal = pyqtSignal(object, object)  # (resultado, df_resultados)
    status_signal = pyqtSignal(str)

    def __init__(self, pasta_raiz, caminho_excel, parent=None):
        super().__init__(parent)
        self.pasta_raiz = pasta_raiz
        self.caminho_excel = caminho_excel
        self.wb = None

    def run(self):
        try:
            self.status_signal.emit("Analisando...")
            self.log_signal.emit("=" * 50)
            self.log_signal.emit(f"INICIANDO ANÁLISE")
            self.log_signal.emit(f"Arquivo Excel: {os.path.basename(self.caminho_excel)}")
            self.log_signal.emit("=" * 50)

            resultado, df_resultados = self.analisar_dados()
            self.finished_signal.emit(resultado, df_resultados)
        except Exception as e:
            self.log_signal.emit(f"✗ Erro durante análise: {str(e)}")
            import traceback
            self.log_signal.emit(traceback.format_exc())
            self.finished_signal.emit(None, None)

    def analisar_dados(self):
        """Lógica de análise real usando openpyxl para carregar"""
        import pandas as pd
        pasta_raiz = self.pasta_raiz
        caminho_excel = self.caminho_excel

        self.log_signal.emit(f"Carregando planilha Excel: {os.path.basename(caminho_excel)}")

        try:
            # Carregar usando openpyxl (suporta .xlsm)
            df_excel, ws = self.carregar_excel_com_openpyxl(caminho_excel)

            if df_excel is None or df_excel.empty:
                self.log_signal.emit("Erro: DataFrame vazio ou não carregado")
                return None, None

            colunas = list(df_excel.columns)
            self.log_signal.emit(f"Colunas encontradas: {', '.join(colunas[:10])}" + ("..." if len(colunas) > 10 else ""))

            # Encontrar coluna de data
            coluna_data = None
            for col in df_excel.columns:
                col_str = str(col).lower()
                if 'data' in col_str or col_str == 'data' or col_str == 'date':
                    coluna_data = col
                    break

            if not coluna_data:
                # Tentar encontrar por padrão comum de datas
                for col in df_excel.columns:
                    if len(df_excel) > 0:
                        sample = str(df_excel[col].iloc[0])
                        if re.search(r'\d{1,2}[/\-]\d{1,2}[/\-]\d{2,4}', sample) or re.search(r'\d{8}', sample):
                            coluna_data = col
                            break

            if not coluna_data:
                coluna_data = df_excel.columns[0]
                self.log_signal.emit(f"Usando primeira coluna como data: '{coluna_data}'")

            self.log_signal.emit(f"Coluna de data identificada: '{coluna_data}'")

            # Converter datas do Excel
            def converter_data_excel(valor):
                if pd.isna(valor):
                    return None
                try:
                    if isinstance(valor, (datetime, pd.Timestamp)):
                        return valor
                    if isinstance(valor, str):
                        valor = valor.strip()
                        formatos = [
                            '%Y-%m-%d %H:%M:%S',
                            '%d/%m/%Y %H:%M:%S',
                            '%d/%m/%Y',
                            '%Y-%m-%d',
                            '%d-%m-%Y',
                            '%d.%m.%Y'
                        ]
                        for formato in formatos:
                            try:
                                return datetime.strptime(valor, formato)
                            except:
                                continue
                    if isinstance(valor, (int, float)):
                        return pd.Timestamp('1899-12-30') + pd.Timedelta(days=valor)
                    return None
                except:
                    return None

            df_excel['Data_Datetime'] = df_excel[coluna_data].apply(converter_data_excel)
            datas_validas = df_excel['Data_Datetime'].notna().sum()
            self.log_signal.emit(f"{datas_validas} datas válidas encontradas no Excel")

            if datas_validas == 0:
                self.log_signal.emit("Nenhuma data válida encontrada no Excel")
                return None, None

        except Exception as e:
            self.log_signal.emit(f"Erro ao carregar Excel: {e}")
            import traceback
            self.log_signal.emit(traceback.format_exc())
            return None, None

        self.log_signal.emit("\nAnalisando subpastas...")

        # Coletar todas as datas das subpastas
        todas_datas_subpastas = {}
        pastas_principais = []

        try:
            pastas_principais = [
                item for item in os.listdir(pasta_raiz)
                if os.path.isdir(os.path.join(pasta_raiz, item))
            ]
            self.log_signal.emit(f"{len(pastas_principais)} pastas principais encontradas")
        except Exception as e:
            self.log_signal.emit(f"Erro ao listar pastas: {e}")
            return None, None

        if not pastas_principais:
            self.log_signal.emit("Nenhuma pasta encontrada na pasta raiz")
            return None, None

        for pasta in pastas_principais:
            caminho_pasta = os.path.join(pasta_raiz, pasta)
            try:
                subpastas = [
                    item for item in os.listdir(caminho_pasta)
                    if os.path.isdir(os.path.join(caminho_pasta, item))
                ]
            except Exception as e:
                self.log_signal.emit(f"Erro ao acessar pasta '{pasta}': {e}")
                continue

            datas_pasta = []
            for subpasta in subpastas:
                data_dt = self.converter_subpasta_para_data(subpasta)
                if data_dt:
                    datas_pasta.append(data_dt)

            if datas_pasta:
                datas_pasta.sort()
                todas_datas_subpastas[pasta] = datas_pasta
                self.log_signal.emit(f"  {pasta}: {len(datas_pasta)} datas válidas")

        total_datas_subpastas = sum(len(datas) for datas in todas_datas_subpastas.values())
        self.log_signal.emit(f"\nTotal de datas encontradas nas subpastas: {total_datas_subpastas}")

        if total_datas_subpastas == 0:
            self.log_signal.emit("Nenhuma data válida encontrada nas subpastas")
            return None, None

        self.log_signal.emit("\nComparando dados entre Excel e subpastas...")
        todos_resultados = []

        # Criar mapa de datas do Excel
        mapa_datas_excel = {}
        for idx, linha in df_excel.iterrows():
            data_excel = linha['Data_Datetime']
            if pd.notna(data_excel):
                data_str = data_excel.strftime('%Y-%m-%d')
                mapa_datas_excel[data_str] = idx

        for pasta, datas_subpastas in todas_datas_subpastas.items():
            # Encontrar coluna correspondente
            coluna_correspondente = None
            for coluna in df_excel.columns:
                if coluna == coluna_data or coluna == 'Data_Datetime':
                    continue
                if str(coluna).strip().upper() == str(pasta).strip().upper():
                    coluna_correspondente = coluna
                    break

            if not coluna_correspondente:
                for coluna in df_excel.columns:
                    if coluna == coluna_data or coluna == 'Data_Datetime':
                        continue
                    col_nome_limpo = str(coluna).strip().upper()
                    pasta_nome_limpo = str(pasta).strip().upper()
                    if pasta_nome_limpo in col_nome_limpo or col_nome_limpo in pasta_nome_limpo:
                        coluna_correspondente = coluna
                        break

            if not coluna_correspondente:
                pasta_sem_espacos = str(pasta).replace(' ', '').upper()
                for coluna in df_excel.columns:
                    if coluna == coluna_data or coluna == 'Data_Datetime':
                        continue
                    coluna_sem_espacos = str(coluna).replace(' ', '').upper()
                    if coluna_sem_espacos == pasta_sem_espacos:
                        coluna_correspondente = coluna
                        break

            if not coluna_correspondente:
                self.log_signal.emit(f"  Pasta '{pasta}': nenhuma coluna correspondente encontrada")
                continue

            self.log_signal.emit(f"  {pasta} -> usando coluna: '{coluna_correspondente}'")

            for data_subpasta in datas_subpastas:
                data_str = data_subpasta.strftime('%Y-%m-%d')
                if data_str in mapa_datas_excel:
                    linha_correspondente = mapa_datas_excel[data_str]
                    valor_celula = df_excel.at[linha_correspondente, coluna_correspondente]
                    esta_vazio = pd.isna(valor_celula) or str(valor_celula).strip() in ['', 'None', 'nan', 'NaN', 'NULL', 'null']

                    resultado = {
                        'pasta_principal': pasta,
                        'data_subpasta': data_subpasta.strftime('%d/%m/%Y'),
                        'data_excel': df_excel.at[linha_correspondente, coluna_data],
                        'coluna_excel': coluna_correspondente,
                        'linha_excel': linha_correspondente + 2,
                        'valor_atual': valor_celula if not pd.isna(valor_celula) else '(vazio)',
                        'esta_vazio': esta_vazio,
                        'acao_recomendada': 'PENDENTE' if esta_vazio else 'MANTER'
                    }
                    todos_resultados.append(resultado)

        if todos_resultados:
            df_resultados = pd.DataFrame(todos_resultados)
            df_resultados = df_resultados.sort_values(['pasta_principal', 'data_subpasta'])
            self.log_signal.emit(f"\n✓ Total de correspondências encontradas: {len(todos_resultados)}")
            return True, df_resultados

        self.log_signal.emit("✗ Nenhuma correspondência encontrada")
        return None, None

    def carregar_excel_com_openpyxl(self, caminho):
        from openpyxl import load_workbook
        try:
            self.log_signal.emit(f"Carregando arquivo com openpyxl: {os.path.basename(caminho)}")
            self.wb = load_workbook(caminho, keep_vba=True, data_only=True)
            ws = None
            for sheet_name in self.wb.sheetnames:
                if 'tarefa' in sheet_name.lower():
                    ws = self.wb[sheet_name]
                    self.log_signal.emit(f"Usando planilha: '{ws.title}'")
                    break
            if ws is None:
                ws = self.wb.active
                self.log_signal.emit(f"Usando planilha ativa: '{ws.title}'")

            # Ler dados para DataFrame
            import pandas as pd
            data = []
            headers = []
            for col in range(1, ws.max_column + 1):
                cell_value = ws.cell(row=1, column=col).value
                headers.append(cell_value if cell_value is not None else f"Coluna_{col}")

            coluna_data_idx = None
            for col in range(1, ws.max_column + 1):
                cell_value = ws.cell(row=1, column=col).value
                if cell_value and 'data' in str(cell_value).lower():
                    coluna_data_idx = col
                    break
            if coluna_data_idx is None:
                coluna_data_idx = 1

            for row in range(2, ws.max_row + 1):
                data_cell = ws.cell(row=row, column=coluna_data_idx).value
                if data_cell is None or (isinstance(data_cell, str) and data_cell.strip() == ''):
                    continue
                row_data = []
                for col in range(1, ws.max_column + 1):
                    row_data.append(ws.cell(row=row, column=col).value)
                data.append(row_data)

            df = pd.DataFrame(data, columns=headers)
            self.log_signal.emit(f"Dados carregados: {len(df)} linhas, {len(df.columns)} colunas")
            return df, ws

        except Exception as e:
            self.log_signal.emit(f"Erro ao carregar com openpyxl: {e}")
            import traceback
            self.log_signal.emit(traceback.format_exc())
            return None, None

    def converter_subpasta_para_data(self, nome_subpasta):
        nome_limpo = str(nome_subpasta).strip()
        digitos = ''.join(filter(str.isdigit, nome_limpo))
        if len(digitos) >= 6:
            data_str = digitos[:6]
            try:
                ano_2digitos = int(data_str[0:2])
                mes = int(data_str[2:4])
                dia = int(data_str[4:6])
                ano = 2000 + ano_2digitos
                if ano > 2050:
                    ano = 1900 + ano_2digitos
                if 1 <= mes <= 12 and 1 <= dia <= 31:
                    try:
                        return datetime(ano, mes, dia)
                    except ValueError:
                        return None
            except (ValueError, TypeError):
                return None
        return None

class AbaAutomacao(QWidget):
    def __init__(self, parent_app):
        super().__init__()
        self.parent_app = parent_app
        self.config = carregar_configuracao()
        self.arquivo_excel_encontrado = None
        self.df_resultados = None
        self.wb = None
        self.thread = None
        self.init_ui()

    def init_ui(self):
        layout = QVBoxLayout(self)

        # Cabeçalho
        cabecalho = QLabel("🤖 AUTOMAÇÃO DE PLANILHA EXCEL")
        cabecalho.setStyleSheet("""
            QLabel {
                font-size: 18px;
                font-weight: bold;
                color: #FF9800;
                padding: 10px;
                background-color: #FFF3E0;
                border-radius: 5px;
            }
        """)
        cabecalho.setAlignment(Qt.AlignCenter)
        layout.addWidget(cabecalho)

        # Configuração de pastas
        grupo_pastas = QGroupBox("Configuração de Pastas")
        grupo_pastas.setStyleSheet("QGroupBox { font-weight: bold; font-size: 14px; }")
        layout_pastas = QGridLayout(grupo_pastas)

        # Linha 0: Pasta raiz (subpastas de datas)
        lbl_pasta_raiz = QLabel("Pasta com subpastas (datas):")
        self.txt_pasta_raiz = QLineEdit()
        self.txt_pasta_raiz.setText(self.config.get("pasta_raiz", ""))
        btn_pasta_raiz = QPushButton("📁 Procurar")
        btn_pasta_raiz.clicked.connect(self.selecionar_pasta_raiz)
        layout_pastas.addWidget(lbl_pasta_raiz, 0, 0)
        layout_pastas.addWidget(self.txt_pasta_raiz, 0, 1)
        layout_pastas.addWidget(btn_pasta_raiz, 0, 2)

        # Linha 1: Pasta do Excel
        lbl_pasta_excel = QLabel("Pasta com planilha Excel:")
        self.txt_pasta_excel = QLineEdit()
        self.txt_pasta_excel.setText(self.config.get("pasta_excel", ""))
        btn_pasta_excel = QPushButton("📁 Procurar")
        btn_pasta_excel.clicked.connect(self.selecionar_pasta_excel)
        layout_pastas.addWidget(lbl_pasta_excel, 1, 0)
        layout_pastas.addWidget(self.txt_pasta_excel, 1, 1)
        layout_pastas.addWidget(btn_pasta_excel, 1, 2)

        # Linha 2: Informação do arquivo encontrado
        self.info_arquivo = QLabel("Arquivo Excel: Não identificado")
        self.info_arquivo.setStyleSheet("color: #757575; font-size: 11px;")
        layout_pastas.addWidget(self.info_arquivo, 2, 0, 1, 3)

        layout.addWidget(grupo_pastas)

        # Botões de ação
        grupo_botoes = QHBoxLayout()
        self.btn_analisar = QPushButton("🔍 Analisar Arquivos")
        self.btn_analisar.setStyleSheet("background-color: #2196F3; color: white; padding: 8px; font-weight: bold;")
        self.btn_analisar.clicked.connect(self.iniciar_analise)

        self.btn_atualizar = QPushButton("📝 Atualizar Excel Original")
        self.btn_atualizar.setStyleSheet("background-color: #4CAF50; color: white; padding: 8px; font-weight: bold;")
        self.btn_atualizar.setEnabled(False)
        self.btn_atualizar.clicked.connect(self.atualizar_excel_original)

        self.btn_procurar_excel = QPushButton("🔎 Procurar Excel")
        self.btn_procurar_excel.setStyleSheet("background-color: #FF9800; color: white; padding: 8px; font-weight: bold;")
        self.btn_procurar_excel.clicked.connect(self.procurar_excel_na_pasta)

        self.btn_limpar = QPushButton("🧹 Limpar Tudo")
        self.btn_limpar.setStyleSheet("background-color: #9E9E9E; color: white; padding: 8px; font-weight: bold;")
        self.btn_limpar.clicked.connect(self.limpar_tudo)

        grupo_botoes.addWidget(self.btn_analisar)
        grupo_botoes.addWidget(self.btn_atualizar)
        grupo_botoes.addWidget(self.btn_procurar_excel)
        grupo_botoes.addWidget(self.btn_limpar)
        grupo_botoes.addStretch()
        layout.addLayout(grupo_botoes)

        # Tabela de resultados
        self.tabela = QTableWidget()
        self.tabela.setColumnCount(6)
        self.tabela.setHorizontalHeaderLabels(["Pasta", "Data", "Linha", "Coluna", "Status", "Ação Necessária"])
        self.tabela.horizontalHeader().setSectionResizeMode(QHeaderView.Stretch)
        self.tabela.setAlternatingRowColors(True)
        layout.addWidget(self.tabela, 1)

        # Log e estatísticas (em duas colunas)
        painel_inferior = QHBoxLayout()

        # Estatísticas
        grupo_stats = QGroupBox("Estatísticas")
        grupo_stats.setStyleSheet("QGroupBox { font-weight: bold; }")
        layout_stats = QVBoxLayout(grupo_stats)
        self.stats_text = QTextEdit()
        self.stats_text.setReadOnly(True)
        self.stats_text.setMaximumHeight(150)
        layout_stats.addWidget(self.stats_text)
        painel_inferior.addWidget(grupo_stats)

        # Log
        grupo_log = QGroupBox("Log de Execução")
        grupo_log.setStyleSheet("QGroupBox { font-weight: bold; }")
        layout_log = QVBoxLayout(grupo_log)
        self.log_text = QTextEdit()
        self.log_text.setReadOnly(True)
        self.log_text.setMaximumHeight(150)
        layout_log.addWidget(self.log_text)
        painel_inferior.addWidget(grupo_log)

        layout.addLayout(painel_inferior)

        # Status bar (dentro da aba)
        self.status_label = QLabel("Pronto para iniciar análise")
        self.status_label.setStyleSheet("background-color: #F5F5F5; padding: 5px; border-radius: 3px;")
        layout.addWidget(self.status_label)

        # Tentar encontrar Excel automaticamente se já houver pasta configurada
        if self.txt_pasta_excel.text() and os.path.exists(self.txt_pasta_excel.text()):
            self.procurar_excel_na_pasta()

    def selecionar_pasta_raiz(self):
        pasta = QFileDialog.getExistingDirectory(self, "Selecionar pasta com subpastas (datas)", self.txt_pasta_raiz.text())
        if pasta:
            self.txt_pasta_raiz.setText(pasta)
            self.config["pasta_raiz"] = pasta
            salvar_configuracao(self.config)

    def selecionar_pasta_excel(self):
        pasta = QFileDialog.getExistingDirectory(self, "Selecionar pasta com planilha Excel", self.txt_pasta_excel.text())
        if pasta:
            self.txt_pasta_excel.setText(pasta)
            self.config["pasta_excel"] = pasta
            salvar_configuracao(self.config)
            self.procurar_excel_na_pasta()

    def procurar_excel_na_pasta(self):
        pasta = self.txt_pasta_excel.text()
        if not pasta or not os.path.exists(pasta):
            QMessageBox.warning(self, "Aviso", "Pasta não existe ou não foi selecionada!")
            return

        self.log("Procurando arquivo Excel na pasta...")
        arquivo = self.procurar_arquivo_excel(pasta)
        if arquivo:
            self.arquivo_excel_encontrado = arquivo
            nome = os.path.basename(arquivo)
            self.info_arquivo.setText(f"Arquivo encontrado: {nome}")
            self.log(f"Arquivo Excel identificado: {nome}")
        else:
            self.arquivo_excel_encontrado = None
            self.info_arquivo.setText("Nenhum arquivo Excel encontrado")
            self.log("Nenhum arquivo Excel encontrado na pasta")
            resposta = QMessageBox.question(self, "Arquivo não encontrado",
                                            "Nenhum arquivo Excel encontrado. Deseja procurar manualmente?",
                                            QMessageBox.Yes | QMessageBox.No)
            if resposta == QMessageBox.Yes:
                self.selecionar_arquivo_excel_manual()

    def procurar_arquivo_excel(self, pasta):
        padroes = [
            "CONTROLE RELATÓRIO.xlsm", "CONTROLE RELATÓRIO.xlsx", "CONTROLE RELATÓRIO.xls",
            "CONTROLE.xlsm", "CONTROLE.xlsx", "CONTROLE.xls",
            "RELATÓRIO.xlsm", "RELATÓRIO.xlsx", "RELATÓRIO.xls"
        ]
        try:
            for arquivo in os.listdir(pasta):
                if arquivo.lower().endswith('.xlsm'):
                    for padrao in padroes:
                        if '*' in padrao:
                            padrao_base = padrao.replace('*', '').replace('.xlsm', '').lower()
                            if padrao_base in arquivo.lower():
                                return os.path.join(pasta, arquivo)
                        else:
                            if arquivo.lower() == padrao.lower():
                                return os.path.join(pasta, arquivo)
            for padrao in padroes:
                if '*' in padrao:
                    for arquivo in os.listdir(pasta):
                        padrao_base = padrao.replace('*', '').split('.')[0].lower()
                        if padrao_base in arquivo.lower() and arquivo.lower().endswith(('.xlsx', '.xls', '.xlsm')):
                            return os.path.join(pasta, arquivo)
                else:
                    caminho = os.path.join(pasta, padrao)
                    if os.path.exists(caminho):
                        return caminho
            for arquivo in os.listdir(pasta):
                if arquivo.lower().endswith(('.xlsm', '.xlsx', '.xls')):
                    return os.path.join(pasta, arquivo)
        except Exception as e:
            self.log(f"Erro ao procurar arquivo Excel: {e}")
        return None

    def selecionar_arquivo_excel_manual(self):
        arquivo, _ = QFileDialog.getOpenFileName(self, "Selecionar arquivo Excel",
                                                  self.txt_pasta_excel.text(),
                                                  "Arquivos Excel (*.xlsx *.xls *.xlsm);;Todos os arquivos (*.*)")
        if arquivo:
            self.arquivo_excel_encontrado = arquivo
            nome = os.path.basename(arquivo)
            pasta = os.path.dirname(arquivo)
            self.txt_pasta_excel.setText(pasta)
            self.config["pasta_excel"] = pasta
            salvar_configuracao(self.config)
            self.info_arquivo.setText(f"Arquivo selecionado: {nome}")
            self.log(f"Arquivo Excel selecionado manualmente: {nome}")

    def iniciar_analise(self):
        if not self.txt_pasta_raiz.text() or not os.path.exists(self.txt_pasta_raiz.text()):
            QMessageBox.critical(self, "Erro", "Pasta com subpastas não encontrada ou não selecionada!")
            return
        if not self.txt_pasta_excel.text() or not os.path.exists(self.txt_pasta_excel.text()):
            QMessageBox.critical(self, "Erro", "Pasta do Excel não encontrada ou não selecionada!")
            return
        if not self.arquivo_excel_encontrado or not os.path.exists(self.arquivo_excel_encontrado):
            QMessageBox.critical(self, "Erro", "Arquivo Excel não identificado! Clique em 'Procurar Excel' primeiro.")
            return

        self.btn_analisar.setEnabled(False)
        self.btn_atualizar.setEnabled(False)
        self.tabela.setRowCount(0)
        self.log_text.clear()
        self.stats_text.clear()

        self.thread = AutomacaoThread(self.txt_pasta_raiz.text(), self.arquivo_excel_encontrado)
        self.thread.log_signal.connect(self.log)
        self.thread.status_signal.connect(self.atualizar_status)
        self.thread.finished_signal.connect(self.on_analise_finished)
        self.thread.start()

    def on_analise_finished(self, resultado, df_resultados):
        self.btn_analisar.setEnabled(True)
        if resultado and df_resultados is not None:
            self.df_resultados = df_resultados
            self.mostrar_resultados()
            self.log("✓ Análise concluída com sucesso!")
        else:
            self.log("✗ Nenhum resultado encontrado na análise")
            self.atualizar_status("Nenhuma correspondência encontrada")
            self.btn_atualizar.setEnabled(False)

    def mostrar_resultados(self):
        if self.df_resultados is None or self.df_resultados.empty:
            return
        self.tabela.setRowCount(len(self.df_resultados))
        for i, row in self.df_resultados.iterrows():
            status = "VAZIO" if row['esta_vazio'] else "PREENCHIDO"
            acao = "PENDENTE" if row['esta_vazio'] else "OK"
            self.tabela.setItem(i, 0, QTableWidgetItem(row['pasta_principal']))
            self.tabela.setItem(i, 1, QTableWidgetItem(row['data_subpasta']))
            self.tabela.setItem(i, 2, QTableWidgetItem(str(row['linha_excel'])))
            self.tabela.setItem(i, 3, QTableWidgetItem(row['coluna_excel']))
            self.tabela.setItem(i, 4, QTableWidgetItem(status))
            self.tabela.setItem(i, 5, QTableWidgetItem(acao))

        total = len(self.df_resultados)
        vazios = self.df_resultados['esta_vazio'].sum()
        preenchidos = total - vazios
        stats = {
            "Total de correspondências": total,
            "Campos já preenchidos": preenchidos,
            "Campos vazios": vazios,
            "Pastas analisadas": self.df_resultados['pasta_principal'].nunique(),
            "Datas analisadas": self.df_resultados['data_subpasta'].nunique(),
            "Arquivo Excel": os.path.basename(self.arquivo_excel_encontrado),
            "Tipo": "Com Macros" if self.arquivo_excel_encontrado.endswith('.xlsm') else "Sem Macros"
        }
        self.atualizar_stats(stats)

        if vazios > 0:
            self.btn_atualizar.setEnabled(True)
            self.atualizar_status(f"{vazios} campo(s) vazio(s) encontrado(s) - Pronto para atualizar")
        else:
            self.btn_atualizar.setEnabled(False)
            self.atualizar_status("Todos os campos já estão preenchidos")

    def atualizar_excel_original(self):
        if self.df_resultados is None or self.df_resultados.empty:
            QMessageBox.warning(self, "Aviso", "Nenhum resultado para atualizar!")
            return

        campos_vazios = self.df_resultados['esta_vazio'].sum()
        if campos_vazios == 0:
            QMessageBox.information(self, "Informação", "Todos os campos já estão preenchidos. Nada para atualizar.")
            return

        resposta = QMessageBox.question(self, "ATENÇÃO - ATUALIZAÇÃO DO ARQUIVO ORIGINAL",
                                        f"Você está prestes a MODIFICAR O ARQUIVO ORIGINAL:\n\n"
                                        f"Arquivo: {os.path.basename(self.arquivo_excel_encontrado)}\n"
                                        f"Local: {os.path.dirname(self.arquivo_excel_encontrado)}\n"
                                        f"Tipo: {'COM MACROS (.xlsm)' if self.arquivo_excel_encontrado.endswith('.xlsm') else 'Sem macros'}\n\n"
                                        f"Serão atualizados {campos_vazios} campo(s) vazio(s) para 'PENDENTE'.\n\n"
                                        f"Esta ação NÃO PODE SER DESFEITA!\n\nDeseja continuar?",
                                        QMessageBox.Yes | QMessageBox.No, QMessageBox.No)
        if resposta != QMessageBox.Yes:
            self.log("Atualização cancelada pelo usuário")
            return

        self.atualizar_status("Atualizando arquivo original...")
        self.log("=" * 50)
        self.log("INICIANDO ATUALIZAÇÃO DO ARQUIVO ORIGINAL")
        self.log(f"Arquivo: {os.path.basename(self.arquivo_excel_encontrado)}")
        self.log(f"Campos para atualizar: {campos_vazios}")
        self.log("=" * 50)

        try:
            alteracoes = self.executar_atualizacao_original()
            if alteracoes > 0:
                self.log(f"✓ ATUALIZAÇÃO CONCLUÍDA: {alteracoes} campo(s) modificado(s)")
                self.atualizar_status(f"Arquivo atualizado: {alteracoes} campo(s)")
                self.df_resultados.loc[self.df_resultados['esta_vazio'], 'esta_vazio'] = False
                self.mostrar_resultados()
                QMessageBox.information(self, "Atualização Concluída",
                                        f"Arquivo original atualizado com sucesso!\n\n"
                                        f"• {alteracoes} campo(s) marcado(s) como 'PENDENTE'\n"
                                        f"• Macros preservadas: {'SIM' if self.arquivo_excel_encontrado.endswith('.xlsm') else 'N/A'}")
            else:
                self.log("✗ Nenhuma alteração foi realizada")
        except Exception as e:
            self.log(f"✗ ERRO ao atualizar arquivo: {str(e)}")
            QMessageBox.critical(self, "Erro na Atualização",
                                 f"Ocorreu um erro durante a atualização:\n\n{str(e)}\n\n"
                                 f"Verifique se o arquivo não está aberto no Excel.")

    def executar_atualizacao_original(self):
        from openpyxl import load_workbook
        from openpyxl.utils import get_column_letter
        caminho = self.arquivo_excel_encontrado

        try:
            if self.wb:
                try:
                    self.wb.close()
                except:
                    pass
            self.wb = load_workbook(caminho, keep_vba=True)

            ws = None
            for sheet in self.wb.sheetnames:
                if 'tarefa' in sheet.lower():
                    ws = self.wb[sheet]
                    break
            if ws is None:
                ws = self.wb.active
            self.log(f"Planilha selecionada: '{ws.title}'")

            # Mapear cabeçalhos
            cabecalhos = {}
            for col in range(1, ws.max_column + 1):
                cell = ws.cell(row=1, column=col)
                if cell.value:
                    cabecalhos[str(cell.value).strip().upper()] = {
                        'coluna': col,
                        'letra': get_column_letter(col),
                        'nome_original': cell.value
                    }

            alteracoes = 0
            registros_vazios = self.df_resultados[self.df_resultados['esta_vazio']]
            self.log(f"Encontrados {len(registros_vazios)} registros vazios para atualizar")

            for _, row in registros_vazios.iterrows():
                coluna_nome = str(row['coluna_excel']).strip().upper()
                linha = int(row['linha_excel'])

                if coluna_nome in cabecalhos:
                    col_info = cabecalhos[coluna_nome]
                    col_idx = col_info['coluna']
                    cell = ws.cell(row=linha, column=col_idx)
                    if cell.value is None or str(cell.value).strip() in ['', 'None', 'nan', 'NaN']:
                        cell.value = "PENDENTE"
                        alteracoes += 1
                        if alteracoes <= 10:
                            self.log(f"  [{alteracoes}] {col_info['letra']}{linha}: 'PENDENTE'")
                        elif alteracoes == 11:
                            self.log(f"  ... continuando atualização ...")

            self.wb.save(caminho)
            self.wb.close()
            self.wb = None
            return alteracoes

        except PermissionError:
            raise Exception("Arquivo está aberto ou sem permissão de escrita. Feche o Excel e tente novamente.")
        except Exception as e:
            if self.wb:
                try:
                    self.wb.close()
                except:
                    pass
                self.wb = None
            raise

    def limpar_tudo(self):
        resposta = QMessageBox.question(self, "Confirmar", "Deseja limpar todos os resultados e logs?",
                                        QMessageBox.Yes | QMessageBox.No)
        if resposta == QMessageBox.Yes:
            self.tabela.setRowCount(0)
            self.stats_text.clear()
            self.log_text.clear()
            self.df_resultados = None
            self.btn_atualizar.setEnabled(False)
            self.atualizar_status("Pronto para iniciar nova análise")
            self.log("Sistema limpo e pronto para nova análise")

    def log(self, mensagem):
        self.log_text.append(f"[{datetime.now().strftime('%H:%M:%S')}] {mensagem}")
        self.log_text.ensureCursorVisible()

    def atualizar_status(self, mensagem):
        self.status_label.setText(mensagem)

    def atualizar_stats(self, stats_dict):
        self.stats_text.clear()
        for chave, valor in stats_dict.items():
            self.stats_text.append(f"• {chave}: {valor}")

class AbaLog(QWidget):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.init_ui()
        # Referência para o app principal (para receber logs)
        self.parent_app = parent

    def init_ui(self):
        layout = QVBoxLayout(self)
        cabecalho = QLabel("📋 LOG DO SISTEMA")
        cabecalho.setStyleSheet("""
            QLabel {
                font-size: 18px;
                font-weight: bold;
                color: #FF5722;
                padding: 10px;
                background-color: #FBE9E7;
                border-radius: 5px;
            }
        """)
        cabecalho.setAlignment(Qt.AlignCenter)
        layout.addWidget(cabecalho)

        self.log_text = QTextEdit()
        self.log_text.setReadOnly(True)
        self.log_text.setStyleSheet("""
            QTextEdit {
                font-family: 'Consolas', monospace;
                font-size: 11px;
                background-color: #1e1e1e;
                color: #d4d4d4;
                border-radius: 5px;
                padding: 5px;
            }
        """)
        layout.addWidget(self.log_text, 1)

        btn_limpar = QPushButton("🧹 Limpar Log")
        btn_limpar.setStyleSheet("background-color: #F44336; color: white; padding: 8px; font-weight: bold;")
        btn_limpar.clicked.connect(self.limpar_log)
        layout.addWidget(btn_limpar)

    def adicionar_log(self, mensagem, nivel="INFO"):
        timestamp = datetime.now().strftime("%H:%M:%S")
        cor = {
            "INFO": "#4CAF50",
            "AVISO": "#FF9800",
            "ERRO": "#F44336",
            "SUCESSO": "#2196F3"
        }.get(nivel, "#FFFFFF")
        html = f'<span style="color:{cor};">[{timestamp}] [{nivel}] {mensagem}</span><br>'
        self.log_text.append(html)
        # Rola para o final
        scrollbar = self.log_text.verticalScrollBar()
        scrollbar.setValue(scrollbar.maximum())

    def limpar_log(self):
        self.log_text.clear()

class App(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("Sistema de Relatórios de Segurança - ENGELMIG")
        self.resize(1400, 900)
        self.lista_imagens_compartilhada = []
        self.init_ui()
        self.aplicar_icone_configurado()
        self.log_message("Aplicação iniciada", "INFO")
        QTimer.singleShot(2000, self.verificar_atualizacao)

    def log_message(self, mensagem, nivel="INFO"):
        """Método global para adicionar mensagens à aba de Log."""
        if hasattr(self, 'aba_log'):
            self.aba_log.adicionar_log(mensagem, nivel)
        # Também imprime no console (útil para debug)
        print(f"[{nivel}] {mensagem}")

    def aplicar_icone_configurado(self):
        config = carregar_configuracao()
        icone_path = config.get("icone_app", "")
        if icone_path and os.path.exists(icone_path):
            self.setWindowIcon(QIcon(icone_path))

    def init_ui(self):
        # Menu
        menu_bar = self.menuBar()
        menu_arquivo = menu_bar.addMenu("Arquivo")
        acao_config = menu_arquivo.addAction("Configurações")
        acao_config.triggered.connect(self.abrir_configuracao)
        menu_arquivo.addSeparator()
        acao_sair = menu_arquivo.addAction("Sair")
        acao_sair.triggered.connect(self.close)
        menu_ajuda = menu_bar.addMenu("Ajuda")
        acao_sobre = menu_ajuda.addAction("ℹ Sobre")
        acao_sobre.triggered.connect(self.mostrar_sobre)

        # Abas
        self.abas = QTabWidget()
        self.aba_relatorio = AbaRelatorio(self)
        self.aba_imagens = AbaImagens(self)
        self.aba_registro = AbaRegistro(self)
        self.aba_automacao = AbaAutomacao(self)
        self.aba_log = AbaLog(self)   # NOVA ABA DE LOG

        self.abas.addTab(self.aba_relatorio, "Novo Relatório")
        self.abas.addTab(self.aba_imagens, "Evidências")
        self.abas.addTab(self.aba_registro, "Registros")
        self.abas.addTab(self.aba_automacao, "Excel")
        self.abas.addTab(self.aba_log, "Log do Sistema")

        self.setCentralWidget(self.abas)
        self.status_bar = self.statusBar()
        self.status_bar.showMessage("Pronto")

    def abrir_configuracao(self):
        dialog = ConfiguracaoDialog(self)
        dialog.exec_()

    def mostrar_sobre(self):
        QMessageBox.about(self, "Sobre", "Sistema de Relatórios - ENGELMIG\nVersão 2.0")

    def verificar_atualizacao(self):
        self.log_message("Verificando atualizações...", "INFO")
        try:
            versao_atual = obter_versao_atual()
            url_api = "https://api.github.com/repos/Bl4ckF/automacao/releases/latest"
            response = requests.get(url_api, timeout=10, headers={'Accept': 'application/vnd.github+json'})
            if response.status_code != 200:
                self.log_message(f"Erro ao buscar release: {response.status_code}", "ERRO")
                return
            data = response.json()
            versao_nova = data.get("tag_name", "").lstrip('v')
            if not versao_nova:
                return
            if self._comparar_versoes(versao_atual, versao_nova) >= 0:
                self.log_message("Programa já está atualizado.", "INFO")
                return
            reply = QMessageBox.question(self, "Atualização", f"Nova versão {versao_nova}. Deseja atualizar?",
                                         QMessageBox.Yes | QMessageBox.No)
            if reply == QMessageBox.Yes:
                asset_url = None
                for asset in data.get("assets", []):
                    if asset["name"].endswith(".exe"):
                        asset_url = asset["browser_download_url"]
                        break
                if asset_url:
                    self.log_message("Baixando atualização...", "INFO")
                    novo_exe_temp = os.path.join(tempfile.gettempdir(), f"app_novo_{versao_nova}.exe")
                    progress = QProgressDialog("Baixando...", "Cancelar", 0, 100, self)
                    progress.setWindowTitle("Atualização")
                    progress.setModal(True)
                    self.download_thread = UpdateDownloadThread(asset_url, novo_exe_temp)
                    self.download_thread.progress.connect(progress.setValue)
                    self.download_thread.finished.connect(lambda path, err: self._atualizacao_baixada(path, err, progress))
                    self.download_thread.start()
                else:
                    self.log_message("Nenhum executável encontrado na release.", "ERRO")
        except Exception as e:
            self.log_message(f"Erro na verificação: {e}", "ERRO")

    def _atualizacao_baixada(self, novo_exe, erro, progress_dialog):
        progress_dialog.close()
        if erro:
            self.log_message(f"Falha no download: {erro}", "ERRO")
            return
        if getattr(sys, 'frozen', False):
            exe_atual = sys.executable
        else:
            exe_atual = os.path.abspath(__file__)
        updater_exe = os.path.join(os.path.dirname(exe_atual), "updater.exe")
        if not os.path.exists(updater_exe):
            self.log_message("updater.exe não encontrado", "ERRO")
            return
        try:
            subprocess.Popen([updater_exe, novo_exe, exe_atual])
            self.log_message("Aplicando atualização...", "SUCESSO")
            QApplication.quit()
            sys.exit(0)
        except Exception as e:
            self.log_message(f"Erro ao iniciar updater: {e}", "ERRO")

    def _comparar_versoes(self, v1, v2):
        def normalize(v):
            return [int(x) for x in v.split('.')]
        try:
            v1_parts = normalize(v1)
            v2_parts = normalize(v2)
            for i in range(max(len(v1_parts), len(v2_parts))):
                a = v1_parts[i] if i < len(v1_parts) else 0
                b = v2_parts[i] if i < len(v2_parts) else 0
                if a != b:
                    return a - b
            return 0
        except:
            return (v1 > v2) - (v1 < v2)

if __name__ == "__main__":
    try:
        init_db()
        sincronizar_erros_padrao()
        print("Banco de dados inicializado")
    except Exception as e:
        print(f"Erro no banco: {e}")
    app = QApplication(sys.argv)
    app.setStyle('Fusion')
    janela = App()
    janela.show()
    sys.exit(app.exec_())